"""
EMO2 - Document Parsers
=======================
Parse various document formats (PDF, Word, Excel) from bytes without disk I/O.
"""

import io
from typing import Optional

import pandas as pd
from docx import Document


def parse_attachment(file_data: bytes, filename: str) -> Optional[str]:
    """
    Parse attachment content without saving to disk.
    
    Args:
        file_data: Raw bytes of the attachment
        filename: Original filename (used to determine file type)
    
    Returns:
        Extracted text content as string, or None if parsing fails
    """
    filename_lower = filename.lower()
    
    try:
        if filename_lower.endswith(('.xlsx', '.xls')):
            return parse_excel(file_data, filename)
        elif filename_lower.endswith('.docx'):
            return parse_docx(file_data)
        elif filename_lower.endswith('.pdf'):
            return parse_pdf(file_data)
        elif filename_lower.endswith(('.txt', '.csv', '.md')):
            return file_data.decode('utf-8', errors='ignore')
        else:
            return f"[Unsupported file type: {filename}]"
    except Exception as e:
        return f"[Error parsing {filename}: {str(e)}]"


def parse_excel(file_data: bytes, filename: str) -> str:
    """Parse Excel file from bytes and return CSV string summary."""
    buffer = io.BytesIO(file_data)
    excel_file = pd.ExcelFile(buffer)
    sheet_names = excel_file.sheet_names
    
    result_parts = [f"📊 Excel File: {filename}"]
    result_parts.append(f"Sheets: {', '.join(sheet_names)}\n")
    
    for sheet_name in sheet_names:
        df = pd.read_excel(buffer, sheet_name=sheet_name)
        result_parts.append(f"--- Sheet: {sheet_name} ---")
        result_parts.append(f"Shape: {df.shape[0]} rows × {df.shape[1]} columns")
        result_parts.append(f"Columns: {', '.join(df.columns.astype(str))}")
        
        if len(df) > 50:
            result_parts.append("(Showing first 50 rows)")
            csv_str = df.head(50).to_csv(index=False)
        else:
            csv_str = df.to_csv(index=False)
        
        result_parts.append(csv_str)
        result_parts.append("")
    
    return "\n".join(result_parts)


def parse_docx(file_data: bytes) -> str:
    """Parse Word document from bytes and extract all text."""
    buffer = io.BytesIO(file_data)
    doc = Document(buffer)
    
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                table_texts.append(row_text)
    
    result_parts = paragraphs
    if table_texts:
        result_parts.append("\n--- Tables ---")
        result_parts.extend(table_texts)
    
    return "\n".join(result_parts)


def parse_pdf(file_data: bytes) -> str:
    """Parse PDF file from bytes and extract text using pypdf."""
    try:
        from pypdf import PdfReader
        
        buffer = io.BytesIO(file_data)
        reader = PdfReader(buffer)
        
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i + 1} ---")
                text_parts.append(page_text)
        
        return "\n".join(text_parts) if text_parts else "[No text extracted from PDF]"
    
    except ImportError:
        return "[PDF parsing skipped - pypdf not installed]"
    except Exception as e:
        return f"[Error parsing PDF: {str(e)}]"
