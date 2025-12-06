"""
Emo - AI Agent with Gmail Integration and Memory
Built with Streamlit, Groq (OpenAI GPT-OSS-120B), Gmail API, and ChromaDB
"""

import os
import io
import json
import uuid
import time
import base64
from typing import Optional, List, Dict
from datetime import datetime, timedelta
from email.mime.text import MIMEText

# Third-party imports
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
import requests
import re
from youtube_transcript_api import YouTubeTranscriptApi

# Google Auth imports
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build

# Data processing imports
import pandas as pd
from docx import Document

# Vector database
import chromadb

# Chat History
from history import (
    load_all_sessions,
    save_session,
    create_new_session,
    generate_title,
    delete_session,
    get_session_messages,
    get_sessions_sorted
)

# Universal Context (4 Pillars)
from state import (
    EmoState,
    initialize_context,
    format_system_prompt,
    refresh_environment,
    update_working_memory,
    load_todo_list,
    get_current_datetime
)

# Session-scoped memory management
from tools import (
    set_current_session_id,
    clear_session_short_term_memory,
)

# Load environment variables
load_dotenv()


# =============================================================================
# STREAMING & RENDERING HELPERS
# =============================================================================

def stream_data(text: str, delay: float = 0.02):
    """
    Generator function for simulated streaming (typing effect).
    
    Args:
        text: The full text to stream
        delay: Time delay between words (default 0.02s)
    
    Yields:
        Words one at a time with a space
    """
    if not isinstance(text, str):
        text = str(text)
    
    for word in text.split(" "):
        yield word + " "
        time.sleep(delay)


def render_message_with_latex(content: str):
    """
    Render a message with proper LaTeX support.
    Converts various LaTeX syntaxes to Streamlit-compatible format.
    Handles: \[...\], \(...\), $...$, $$...$$
    """
    import re
    
    if not content:
        return
    
    # Fix setext-style headers: text followed by line of === or ---
    # These cause random text to become huge headers
    # We escape by adding a zero-width space after the text
    content = re.sub(
        r'^([^\n#*`\->][^\n]*)\n(={3,}|-{3,})$',
        r'\1\n\n\2',  # Add extra newline to break the header pattern
        content,
        flags=re.MULTILINE
    )
    
    # Convert \[...\] (display math) to $$ blocks
    # These need to be on their own lines in Streamlit
    content = re.sub(
        r'\\\[(.*?)\\\]',
        lambda m: f'\n\n$$\n{m.group(1).strip()}\n$$\n\n',
        content,
        flags=re.DOTALL
    )
    
    # Convert \(...\) (inline math) to $...$
    content = re.sub(
        r'\\\((.*?)\\\)',
        lambda m: f'${m.group(1).strip()}$',
        content,
        flags=re.DOTALL
    )
    
    # Fix standalone $$...$$ on single line - split them properly
    def fix_block_latex(match):
        latex = match.group(1).strip()
        return f'\n\n$$\n{latex}\n$$\n\n'
    
    # Match $$ that are on their own line (block equations)
    content = re.sub(
        r'(?:^|\n)\s*\$\$([^$]+)\$\$\s*(?:\n|$)',
        fix_block_latex,
        content
    )
    
    # Clean up extra newlines
    content = re.sub(r'\n{4,}', '\n\n\n', content)
    
    # Render with markdown
    st.markdown(content, unsafe_allow_html=True)


def process_latex_content(text: str) -> str:
    """
    Process text to convert LaTeX syntax for display.
    Returns the processed string (doesn't render it).
    Handles: \\[...\\], \\(...\\), $...$, $$...$$
    """
    import re
    
    if not text:
        return text
    
    # Convert \[...\] (display math) to $$ blocks
    text = re.sub(
        r'\\\[(.+?)\\\]',
        lambda m: f'$${m.group(1).strip()}$$',
        text,
        flags=re.DOTALL
    )
    
    # Convert \(...\) (inline math) to $...$
    text = re.sub(
        r'\\\((.+?)\\\)',
        lambda m: f'${m.group(1).strip()}$',
        text,
        flags=re.DOTALL
    )
    
    return text

# Gmail API scopes - readonly access to Gmail
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']

# File paths
CREDENTIALS_FILE = 'credentials.json'
TOKEN_FILE = 'token.json'
TODO_FILE = 'todo.json'


# =============================================================================
# TASK 1: SETUP & AUTHENTICATION
# =============================================================================

def authenticate_gmail():
    """
    Authenticate with Gmail API using OAuth 2.0.
    
    - Loads credentials from credentials.json
    - Saves/loads user session to token.json to avoid re-authentication
    - Returns an authenticated Gmail API service object
    """
    creds = None
    
    # Check if token.json exists (previously saved credentials)
    if os.path.exists(TOKEN_FILE):
        creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
    
    # If no valid credentials available, authenticate
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            # Refresh expired credentials
            creds.refresh(Request())
        else:
            # Run OAuth flow for new authentication
            if not os.path.exists(CREDENTIALS_FILE):
                raise FileNotFoundError(
                    f"'{CREDENTIALS_FILE}' not found. Please download it from "
                    "Google Cloud Console and place it in the project root."
                )
            
            flow = InstalledAppFlow.from_client_secrets_file(
                CREDENTIALS_FILE, SCOPES
            )
            creds = flow.run_local_server(port=0)
        
        # Save credentials for future runs
        with open(TOKEN_FILE, 'w') as token:
            token.write(creds.to_json())
    
    # Build and return the Gmail API service
    service = build('gmail', 'v1', credentials=creds)
    return service


def get_gmail_service():
    """
    Get or create Gmail service instance.
    Uses Streamlit session state to cache the service.
    """
    if 'gmail_service' not in st.session_state:
        st.session_state.gmail_service = authenticate_gmail()
    return st.session_state.gmail_service


def disconnect_gmail():
    """
    Disconnect Gmail by clearing cached service and optionally removing token.
    Returns True if successful.
    """
    try:
        # Clear cached service
        if 'gmail_service' in st.session_state:
            del st.session_state.gmail_service
        
        # Clear authentication state
        st.session_state.gmail_authenticated = False
        
        # Remove token file to force re-authentication
        token_path = os.path.join(os.path.dirname(__file__), 'token.json')
        if os.path.exists(token_path):
            os.remove(token_path)
            print("[Gmail] Token removed")
        
        return True
    except Exception as e:
        print(f"[Gmail] Disconnect error: {e}")
        return False


def reconnect_gmail():
    """
    Force reconnect to Gmail by clearing cache and re-authenticating.
    Returns True if successful, error message if failed.
    """
    try:
        # Clear existing service
        if 'gmail_service' in st.session_state:
            del st.session_state.gmail_service
        
        # Re-authenticate
        service = authenticate_gmail()
        if service:
            st.session_state.gmail_service = service
            st.session_state.gmail_authenticated = True
            return True
        return "Authentication failed"
    except Exception as e:
        st.session_state.gmail_authenticated = False
        return str(e)


def test_gmail_connection():
    """
    Test Gmail connection by fetching a single recent email.
    Returns (success: bool, message: str)
    """
    try:
        service = get_gmail_service()
        results = service.users().messages().list(
            userId='me',
            q='newer_than:1d',
            maxResults=1
        ).execute()
        
        count = len(results.get('messages', []))
        return True, f"Connected! Found {count} email(s) from today."
    except Exception as e:
        return False, f"Error: {str(e)[:50]}"


# Configure Groq API (OpenAI-compatible)
GROQ_API_KEY = os.getenv('GROQ_API_KEY')

# Multi-model architecture for better performance
MODELS = {
    'fast': 'llama-3.1-8b-instant',           # Fast responses, simple tasks (8B)
    'balanced': 'qwen/qwen3-32b',              # Good balance - 32B model
    'tools': 'meta-llama/llama-4-scout-17b-16e-instruct',  # Scout model for tools
}

# Default model for general use
GROQ_MODEL = MODELS['balanced']

def get_groq_client():
    """Get or create Groq OpenAI-compatible client."""
    if 'groq_client' not in st.session_state:
        st.session_state.groq_client = OpenAI(
            api_key=GROQ_API_KEY,
            base_url="https://api.groq.com/openai/v1"
        )
    return st.session_state.groq_client


def classify_query(message: str) -> tuple[str, list[str]]:
    """
    Fast keyword-based router to determine query type and needed tools.
    No AI call needed - just pattern matching.
    
    Returns:
        (query_type, list_of_tools_to_run)
    """
    msg_lower = message.lower()
    
    # Email patterns
    email_keywords = ['email', 'mail', 'gmail', 'inbox', 'message from', 'sent by']
    attachment_keywords = ['attachment', 'attached', 'file', 'document', 'pdf', 'doc']
    
    # Todo patterns  
    todo_keywords = ['todo', 'task', 'remind', 'reminder', 'add to', 'schedule']
    todo_check = ['what are my', 'show me my', 'list my', 'my tasks', 'my todos']
    
    # Web patterns
    web_keywords = ['website', 'webpage', 'url', 'http', 'www', 'read this', 'check this link']
    youtube_keywords = ['youtube', 'video', 'watch', 'yt.com', 'youtu.be']
    news_keywords = ['news', 'headlines', 'what\'s happening', 'current events']
    
    # Quiz patterns
    quiz_keywords = ['quiz', 'test me', 'test my', 'quizz', 'generate quiz', 'make quiz', 'create quiz']
    
    # Memory patterns
    memory_keywords = ['remember', 'recall', 'what did', 'previously', 'earlier', 'you told me']
    save_keywords = ['save this', 'remember this', 'note this', 'store this']
    
    # Simple chat (no tools needed)
    greeting_keywords = ['hi', 'hello', 'hey', 'how are you', 'what\'s up', 'thanks', 'thank you']
    
    tools_needed = []
    query_type = 'chat'  # default
    
    # Check for greetings first (no tools)
    if any(msg_lower.strip().startswith(g) or msg_lower.strip() == g for g in greeting_keywords):
        return ('simple_chat', [])
    
    # Check for quiz generation (should use tools model for JSON generation)
    if any(kw in msg_lower for kw in quiz_keywords):
        query_type = 'quiz'
        # Don't add tool here - let AI decide to call generate_quiz with proper JSON
        return (query_type, [])
    
    # Check for email queries
    # First check if user wants to ANALYZE/READ attachments (not just find emails with attachments)
    analyze_attachment_keywords = ['analyze', 'summarize', 'read the', 'show the', 'what\'s in the', 'open the', 'check the']
    if any(kw in msg_lower for kw in attachment_keywords):
        if any(kw in msg_lower for kw in analyze_attachment_keywords) or 'all' in msg_lower:
            # User wants to analyze/read attachment content
            tools_needed.append('analyze_attachment')
            query_type = 'attachment'
        elif any(kw in msg_lower for kw in email_keywords):
            # User asking about emails with attachments
            tools_needed.append('quick_gmail_search')
            query_type = 'email'
    elif any(kw in msg_lower for kw in email_keywords):
        tools_needed.append('quick_gmail_search')
        query_type = 'email'
    
    # Check for todos
    if any(kw in msg_lower for kw in todo_keywords):
        if any(kw in msg_lower for kw in ['add', 'create', 'new', 'set']):
            tools_needed.append('add_todo')
        elif any(kw in msg_lower for kw in ['done', 'complete', 'finish', 'mark']):
            tools_needed.append('complete_todo')
        else:
            tools_needed.append('get_todos')
        query_type = 'todo'
    elif any(kw in msg_lower for kw in todo_check):
        tools_needed.append('get_todos')
        query_type = 'todo'
    
    # Check for web
    if any(kw in msg_lower for kw in youtube_keywords):
        tools_needed.append('watch_youtube')
        query_type = 'web'
    elif any(kw in msg_lower for kw in web_keywords) or 'http' in msg_lower:
        tools_needed.append('read_web_page')
        query_type = 'web'
    elif any(kw in msg_lower for kw in news_keywords):
        tools_needed.append('get_news_headlines')
        query_type = 'web'
    
    # Check for memory
    if any(kw in msg_lower for kw in save_keywords):
        tools_needed.append('save_long_term_memory')
        query_type = 'memory'
    elif any(kw in msg_lower for kw in memory_keywords):
        tools_needed.append('search_memory')
        query_type = 'memory'
    
    # If no tools detected but seems like a question, might need memory search
    if not tools_needed and '?' in message:
        # Could benefit from memory context, but don't force tool call
        query_type = 'question'
    
    return (query_type, tools_needed)


def select_model_for_task(query_type: str, has_tools: bool) -> str:
    """Select the best model based on task complexity."""
    if query_type == 'simple_chat':
        return MODELS['fast']
    elif query_type == 'quiz':
        # Quiz generation needs the tools model for proper JSON output
        return MODELS['tools']
    elif has_tools:
        return MODELS['tools']
    else:
        return MODELS['balanced']


# =============================================================================
# UNIVERSAL CONTEXT: TO-DO LIST MANAGER
# =============================================================================

class TaskManager:
    """
    Task manager with deadline support and smart reminders.
    Persists tasks to a local JSON file.
    """
    
    def __init__(self, filepath: str = TODO_FILE):
        self.filepath = filepath
        self.tasks = self._load_tasks()
    
    def _load_tasks(self) -> List[Dict]:
        """Load tasks from JSON file."""
        if os.path.exists(self.filepath):
            try:
                with open(self.filepath, 'r') as f:
                    return json.load(f)
            except (json.JSONDecodeError, IOError):
                return []
        return []
    
    def _save_tasks(self):
        """Save tasks to JSON file."""
        with open(self.filepath, 'w') as f:
            json.dump(self.tasks, f, indent=2)
    
    def _parse_deadline(self, task_text: str) -> Optional[datetime]:
        """
        Extract deadline from task text. Language-agnostic - focuses on time patterns.
        Supports formats like:
        - "meeting at 11pm", "call at 3:30 pm", "at 11:00 pm"
        - "Dec 6, 2025 at 2 PM", "December 6, 2025, from 2:00 PM"
        - "2025-12-03 23:00"
        - "tomorrow 5pm", "Friday 2pm"
        - Vietnamese: "lúc 11 giờ", "11h", "11h30"
        """
        import re
        from datetime import timedelta
        
        text_lower = task_text.lower()
        now = datetime.now()
        
        # Try to extract date first
        target_date = now.date()
        
        # Pattern: "Dec 6, 2025" or "December 6, 2025"
        month_names = {
            'jan': 1, 'january': 1, 'feb': 2, 'february': 2, 'mar': 3, 'march': 3,
            'apr': 4, 'april': 4, 'may': 5, 'jun': 6, 'june': 6,
            'jul': 7, 'july': 7, 'aug': 8, 'august': 8, 'sep': 9, 'september': 9,
            'oct': 10, 'october': 10, 'nov': 11, 'november': 11, 'dec': 12, 'december': 12
        }
        
        date_pattern = r'(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|june?|july?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s+(\d{1,2})(?:,?\s*(\d{4}))?'
        date_match = re.search(date_pattern, text_lower)
        if date_match:
            month = month_names.get(date_match.group(1)[:3], now.month)
            day = int(date_match.group(2))
            year = int(date_match.group(3)) if date_match.group(3) else now.year
            try:
                from datetime import date
                target_date = date(year, month, day)
            except:
                pass
        
        # Pattern: ISO date "2025-12-03"
        iso_date = re.search(r'(\d{4})-(\d{2})-(\d{2})', task_text)
        if iso_date:
            try:
                from datetime import date
                target_date = date(int(iso_date.group(1)), int(iso_date.group(2)), int(iso_date.group(3)))
            except:
                pass
        
        # Handle relative dates
        if 'tomorrow' in text_lower or 'ngày mai' in text_lower:
            target_date = (now + timedelta(days=1)).date()
        
        # Now extract time
        hour, minute = None, 0
        
        # Pattern: "HH:MM AM/PM" or "H:MM AM/PM" or "H AM/PM"
        time_patterns = [
            r'(\d{1,2}):(\d{2})\s*(am|pm|a\.m\.|p\.m\.)?',  # 11:30 pm, 2:00 PM
            r'(?:at|by|@|lúc|vào)\s*(\d{1,2})(?::(\d{2}))?\s*(am|pm|a\.m\.|p\.m\.|giờ|h)?',  # at 11pm, lúc 11 giờ
            r'(\d{1,2})\s*(am|pm|a\.m\.|p\.m\.)',  # 11pm, 2PM
            r'(\d{1,2})h(\d{2})?',  # 11h, 11h30 (Vietnamese style)
            r'from\s*(\d{1,2}):(\d{2})\s*(am|pm|a\.m\.|p\.m\.)?',  # from 2:00 PM
        ]
        
        for pattern in time_patterns:
            match = re.search(pattern, text_lower)
            if match:
                groups = match.groups()
                hour = int(groups[0])
                
                # Get minutes if available
                if len(groups) > 1 and groups[1] and groups[1].isdigit():
                    minute = int(groups[1])
                
                # Get AM/PM if available
                ampm = None
                for g in groups:
                    if g and isinstance(g, str) and g.lower().replace('.', '') in ['am', 'pm']:
                        ampm = g.lower().replace('.', '')
                        break
                
                # Apply AM/PM conversion
                if ampm == 'pm' and hour != 12:
                    hour += 12
                elif ampm == 'am' and hour == 12:
                    hour = 0
                # If no AM/PM and hour <= 12 and it's evening context, assume PM
                elif ampm is None and hour <= 12 and hour >= 1:
                    # Assume PM for hours 1-11 if no indicator (common for meetings)
                    if hour < 7:  # 1-6 more likely PM
                        hour += 12
                
                break
        
        if hour is not None:
            try:
                deadline = datetime.combine(target_date, datetime.min.time().replace(hour=hour, minute=minute))
                # If deadline is in the past and no specific date was given, move to tomorrow
                if deadline < now and not date_match and not iso_date:
                    deadline += timedelta(days=1)
                return deadline
            except:
                pass
        
        return None
    
    def add_task(self, task: str) -> Dict:
        """Add a new task with optional deadline detection."""
        deadline = self._parse_deadline(task)
        
        new_task = {
            'id': str(uuid.uuid4()),
            'task': task,
            'status': 'pending',
            'created_at': datetime.now().isoformat(),
            'deadline': deadline.isoformat() if deadline else None
        }
        self.tasks.append(new_task)
        self._save_tasks()
        return new_task
    
    def get_pending_tasks(self) -> List[Dict]:
        """Return all pending tasks."""
        return [t for t in self.tasks if t['status'] == 'pending']
    
    def get_all_tasks(self) -> List[Dict]:
        """Return all tasks."""
        return self.tasks
    
    def get_upcoming_deadlines(self, within_minutes: int = 60) -> List[Dict]:
        """
        Get tasks with deadlines coming up within specified minutes.
        Returns tasks sorted by deadline (soonest first).
        """
        now = datetime.now()
        upcoming = []
        
        for task in self.get_pending_tasks():
            if task.get('deadline'):
                try:
                    deadline = datetime.fromisoformat(task['deadline'])
                    minutes_until = (deadline - now).total_seconds() / 60
                    
                    # Include if deadline is within range and not passed
                    if 0 < minutes_until <= within_minutes:
                        upcoming.append({
                            **task,
                            'minutes_until': int(minutes_until),
                            'deadline_dt': deadline
                        })
                except:
                    pass
        
        # Sort by deadline (soonest first)
        upcoming.sort(key=lambda x: x['deadline_dt'])
        return upcoming
    
    def get_overdue_tasks(self) -> List[Dict]:
        """Get tasks that are past their deadline."""
        now = datetime.now()
        overdue = []
        
        for task in self.get_pending_tasks():
            if task.get('deadline'):
                try:
                    deadline = datetime.fromisoformat(task['deadline'])
                    if deadline < now:
                        minutes_overdue = int((now - deadline).total_seconds() / 60)
                        overdue.append({
                            **task,
                            'minutes_overdue': minutes_overdue,
                            'deadline_dt': deadline
                        })
                except:
                    pass
        
        return overdue
    
    def complete_task_by_index(self, index: int) -> bool:
        """Mark a task as done by its 1-based index in pending tasks."""
        pending = self.get_pending_tasks()
        if 1 <= index <= len(pending):
            task_id = pending[index - 1]['id']
            return self.complete_task_by_id(task_id)
        return False
    
    def complete_task_by_id(self, task_id: str) -> bool:
        """Mark a task as done by its ID."""
        for task in self.tasks:
            if task['id'] == task_id:
                task['status'] = 'done'
                self._save_tasks()
                return True
        return False
    
    def delete_completed(self):
        """Remove all completed tasks."""
        self.tasks = [t for t in self.tasks if t['status'] != 'done']
        self._save_tasks()
    
    def reload(self):
        """Reload tasks from file and backfill missing deadlines."""
        self.tasks = self._load_tasks()
        self._backfill_deadlines()
    
    def _backfill_deadlines(self):
        """Parse deadlines for existing tasks that don't have them."""
        updated = False
        for task in self.tasks:
            if task.get('deadline') is None and task.get('status') == 'pending':
                deadline = self._parse_deadline(task['task'])
                if deadline:
                    task['deadline'] = deadline.isoformat()
                    updated = True
        if updated:
            self._save_tasks()


def get_task_manager() -> TaskManager:
    """
    Get or create TaskManager instance.
    Always creates fresh to ensure new methods are available.
    """
    # Always create fresh instance to pick up new methods
    st.session_state.task_manager = TaskManager()
    return st.session_state.task_manager


def get_smart_reminders() -> str:
    """
    Check for upcoming deadlines and overdue tasks.
    Returns a context string for the AI to naturally mention.
    """
    manager = get_task_manager()
    reminders = []
    
    # Check overdue tasks
    overdue = manager.get_overdue_tasks()
    for task in overdue:
        mins = task['minutes_overdue']
        if mins < 60:
            time_str = f"{mins} minutes ago"
        else:
            hours = mins // 60
            time_str = f"{hours} hour{'s' if hours > 1 else ''} ago"
        reminders.append(f"OVERDUE ({time_str}): {task['task']}")
    
    # Check upcoming deadlines (within 2 hours)
    upcoming = manager.get_upcoming_deadlines(within_minutes=120)
    for task in upcoming:
        mins = task['minutes_until']
        deadline_time = task['deadline_dt'].strftime("%I:%M %p").lstrip('0')
        
        if mins <= 15:
            urgency = "URGENT"
        elif mins <= 30:
            urgency = "SOON"
        else:
            urgency = "UPCOMING"
        
        if mins < 60:
            time_str = f"in {mins} min"
        else:
            hours = mins // 60
            remaining_mins = mins % 60
            if remaining_mins > 0:
                time_str = f"in {hours}h {remaining_mins}m"
            else:
                time_str = f"in {hours} hour{'s' if hours > 1 else ''}"
        
        reminders.append(f"{urgency} ({time_str}, at {deadline_time}): {task['task']}")
    
    if not reminders:
        return ""
    
    return "[DEADLINE ALERTS - mention naturally if relevant]\n" + "\n".join(reminders)


# =============================================================================
# TO-DO TOOL FUNCTIONS (For Gemini)
# =============================================================================

def add_todo(task: str) -> str:
    """
    Add a new task to the to-do list.
    
    Args:
        task: The task description to add.
    
    Returns:
        Confirmation message with the added task.
    """
    manager = get_task_manager()
    new_task = manager.add_task(task)
    pending_count = len(manager.get_pending_tasks())
    
    # Flag that todos changed (for UI refresh)
    st.session_state.todo_changed = True
    
    return f"✅ Added task: '{task}'\n📋 You now have {pending_count} pending task(s)."


def get_todos() -> str:
    """
    Get the list of current pending tasks.
    
    Returns:
        A formatted string of all pending tasks, or a message if none exist.
    """
    manager = get_task_manager()
    pending = manager.get_pending_tasks()
    
    if not pending:
        return "📋 Your to-do list is empty. No pending tasks!"
    
    task_lines = ["📋 **Current To-Do List:**"]
    for i, task in enumerate(pending, 1):
        created = task.get('created_at', 'Unknown')[:10]  # Get date portion
        task_lines.append(f"{i}. {task['task']} (added: {created})")
    
    task_lines.append(f"\n**Total: {len(pending)} pending task(s)**")
    return "\n".join(task_lines)


def complete_todo(task_index: int) -> str:
    """
    Mark a task as done based on its index (1-based) in the pending list.
    
    Args:
        task_index: The 1-based index of the task to complete.
    
    Returns:
        Confirmation message or error if task not found.
    """
    manager = get_task_manager()
    pending = manager.get_pending_tasks()
    
    if not pending:
        return "📋 Your to-do list is empty. Nothing to complete!"
    
    if task_index < 1 or task_index > len(pending):
        return f"❌ Invalid task index. Please choose a number between 1 and {len(pending)}."
    
    task_name = pending[task_index - 1]['task']
    success = manager.complete_task_by_index(task_index)
    
    if success:
        remaining = len(manager.get_pending_tasks())
        # Flag that todos changed (for UI refresh)
        st.session_state.todo_changed = True
        return f"✅ Completed: '{task_name}'\n📋 {remaining} task(s) remaining."
    else:
        return f"❌ Failed to complete task at index {task_index}."


# =============================================================================
# INTERACTIVE QUIZ GENERATOR
# =============================================================================

def generate_quiz(quiz_json: str) -> str:
    """
    Generate an interactive quiz that will be displayed in the UI.
    The AI should call this with a JSON string containing the quiz structure.
    
    Args:
        quiz_json: A JSON string with the following structure:
            {
                "title": "Quiz Title",
                "description": "Optional description",
                "questions": [
                    {
                        "id": 1,
                        "type": "multiple_choice",
                        "question": "What is 2+2?",
                        "options": ["3", "4", "5", "6"],
                        "correct": 1,  // 0-indexed
                        "explanation": "2+2 equals 4"
                    },
                    {
                        "id": 2,
                        "type": "short_answer",
                        "question": "What is the capital of France?",
                        "correct": "Paris",
                        "explanation": "Paris is the capital city of France"
                    },
                    {
                        "id": 3,
                        "type": "true_false",
                        "question": "The Earth is flat.",
                        "correct": false,
                        "explanation": "The Earth is roughly spherical"
                    }
                ]
            }
    
    Returns:
        Confirmation that quiz was created and stored for display.
    """
    try:
        quiz_data = json.loads(quiz_json)
        
        # Validate required fields
        if 'questions' not in quiz_data:
            return "Error: Quiz must contain 'questions' array"
        
        if not quiz_data['questions']:
            return "Error: Quiz must have at least one question"
        
        # Add default title if missing
        if 'title' not in quiz_data:
            quiz_data['title'] = "Quiz"
        
        # Flexible type detection - AI can use any type it wants
        for i, q in enumerate(quiz_data['questions']):
            if 'id' not in q:
                q['id'] = i + 1
            
            # Auto-detect type if not specified
            if 'type' not in q:
                if 'options' in q:
                    q['type'] = 'multiple_choice'
                elif 'pairs' in q or 'matches' in q:
                    q['type'] = 'matching'
                elif 'items' in q or 'order' in q:
                    q['type'] = 'ordering'
                elif 'blanks' in q or '___' in q.get('question', ''):
                    q['type'] = 'fill_blank'
                elif isinstance(q.get('correct'), bool):
                    q['type'] = 'true_false'
                else:
                    q['type'] = 'short_answer'
        
        # Store quiz in session state for UI rendering
        quiz_id = f"quiz_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}"
        quiz_data['quiz_id'] = quiz_id
        quiz_data['created_at'] = datetime.now().isoformat()
        quiz_data['user_answers'] = {}
        quiz_data['submitted'] = False
        quiz_data['score'] = None
        
        # Store in session state
        if 'active_quizzes' not in st.session_state:
            st.session_state.active_quizzes = {}
        
        st.session_state.active_quizzes[quiz_id] = quiz_data
        st.session_state.current_quiz_id = quiz_id
        
        # Also store quiz JSON in a separate dict for persistence across reruns
        if 'quiz_data_store' not in st.session_state:
            st.session_state.quiz_data_store = {}
        st.session_state.quiz_data_store[quiz_id] = quiz_data
        
        num_questions = len(quiz_data['questions'])
        return f"QUIZ_CREATED:{quiz_id}|{quiz_data['title']}|{num_questions} questions"
        
    except json.JSONDecodeError as e:
        return f"Error: Invalid JSON format - {str(e)}"
    except Exception as e:
        return f"Error creating quiz: {str(e)}"


def check_answer(user_ans, correct_ans, q_type: str, question: dict) -> bool:
    """
    Check if user answer is correct.
    Only supports auto-gradable types: multiple_choice and true_false.
    """
    if user_ans is None:
        return False
    
    if q_type == 'multiple_choice':
        return user_ans == correct_ans
    
    elif q_type == 'true_false':
        return user_ans == correct_ans
    
    else:
        # Fallback: direct comparison
        return user_ans == correct_ans


def render_question_input(quiz: dict, question: dict, q_key: str, q_type: str):
    """
    Render the appropriate input widget for a question type.
    Only supports auto-gradable types (no text input needed).
    Supports LaTeX in options for math quizzes.
    """
    q_id = str(question['id'])
    
    if q_type == 'multiple_choice':
        options = question.get('options', [])
        if not options:
            st.warning("No options provided for this question")
            return
        
        current_val = quiz['user_answers'].get(q_id)
        
        # Render options with LaTeX support
        # Use selectbox for non-LaTeX, custom buttons for LaTeX
        has_latex = any('$' in str(opt) or '\\' in str(opt) for opt in options)
        
        if has_latex:
            # LaTeX options - render each with markdown + button
            for j, opt in enumerate(options):
                opt_text = process_latex_content(opt)
                is_selected = current_val == j
                
                # Create a container for each option
                opt_container = st.container()
                with opt_container:
                    cols = st.columns([0.08, 0.82, 0.1])
                    with cols[0]:
                        # Selection indicator
                        if is_selected:
                            st.markdown("**✓**")
                        else:
                            st.markdown(f"**{chr(65+j)}.**")
                    with cols[1]:
                        # Render LaTeX content
                        st.markdown(opt_text)
                    with cols[2]:
                        # Select button
                        if st.button("Select" if not is_selected else "✓", 
                                   key=f"{q_key}_opt{j}",
                                   type="primary" if is_selected else "secondary"):
                            quiz['user_answers'][q_id] = j
        else:
            # Simple options - use radio for better UX
            option_labels = [f"{chr(65+j)}. {opt}" for j, opt in enumerate(options)]
            current_idx = current_val if current_val is not None else None
            
            selected = st.radio(
                "Select answer:",
                options=range(len(options)),
                format_func=lambda x: option_labels[x],
                index=current_idx,
                key=q_key,
                label_visibility="collapsed"
            )
            
            if selected is not None:
                quiz['user_answers'][q_id] = selected
    
    elif q_type == 'true_false':
        current_val = quiz['user_answers'].get(q_id)
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("True", key=f"{q_key}_true", 
                        type="primary" if current_val == True else "secondary",
                        use_container_width=True):
                quiz['user_answers'][q_id] = True
        with col2:
            if st.button("False", key=f"{q_key}_false",
                        type="primary" if current_val == False else "secondary",
                        use_container_width=True):
                quiz['user_answers'][q_id] = False
    
    else:
        # Fallback: treat as multiple choice with LaTeX support if options exist
        options = question.get('options', [])
        if options:
            current_val = quiz['user_answers'].get(q_id)
            has_latex = any('$' in str(opt) or '\\' in str(opt) for opt in options)
            
            if has_latex:
                for j, opt in enumerate(options):
                    opt_text = process_latex_content(opt)
                    is_selected = current_val == j
                    cols = st.columns([0.08, 0.82, 0.1])
                    with cols[0]:
                        st.markdown(f"**{chr(65+j)}.**" if not is_selected else "**✓**")
                    with cols[1]:
                        st.markdown(opt_text)
                    with cols[2]:
                        if st.button("Select" if not is_selected else "✓",
                                   key=f"{q_key}_fallback_opt{j}",
                                   type="primary" if is_selected else "secondary"):
                            quiz['user_answers'][q_id] = j
            else:
                option_labels = [f"{chr(65+j)}. {opt}" for j, opt in enumerate(options)]
                current_idx = current_val if current_val is not None else None
                selected = st.radio("Select:", options=range(len(options)),
                                   format_func=lambda x: option_labels[x],
                                   index=current_idx, key=f"{q_key}_fallback",
                                   label_visibility="collapsed")
                if selected is not None:
                    quiz['user_answers'][q_id] = selected
        else:
            st.info("This question type is not supported for auto-grading")


@st.fragment
def render_quiz(quiz_id: str):
    """
    Render an interactive quiz in the Streamlit UI.
    Uses @st.fragment for partial reruns - only quiz updates, not whole page!
    """
    # Try to get quiz from multiple sources
    quiz = None
    
    if 'active_quizzes' in st.session_state:
        quiz = st.session_state.active_quizzes.get(quiz_id)
    
    # Fallback to data store
    if not quiz and 'quiz_data_store' in st.session_state:
        quiz = st.session_state.quiz_data_store.get(quiz_id)
        if quiz:
            # Restore to active quizzes
            if 'active_quizzes' not in st.session_state:
                st.session_state.active_quizzes = {}
            st.session_state.active_quizzes[quiz_id] = quiz
    
    if not quiz:
        st.info(f"📝 Quiz was generated but needs to be recreated. Ask me to generate another quiz!")
        return
    
    # Quiz container with styling
    with st.container():
        # Header
        st.markdown(f"### 📝 {quiz.get('title', 'Quiz')}")
        if quiz.get('description'):
            st.caption(quiz['description'])
        
        # Progress bar
        total_q = len(quiz['questions'])
        answered = len(quiz.get('user_answers', {}))
        st.progress(answered / total_q if total_q > 0 else 0, 
                   text=f"Progress: {answered}/{total_q} answered")
        
        st.divider()
        
        # Question type icons (only auto-gradable types)
        type_icons = {
            "multiple_choice": "🔘", 
            "true_false": "✓✗",
        }
        
        # Render each question
        for i, question in enumerate(quiz['questions']):
            q_id = question['id']
            q_key = f"{quiz_id}_q{q_id}"
            q_type = question.get('type', 'multiple_choice')
            
            # Question header with icon - LaTeX support
            q_icon = type_icons.get(q_type, "🔘")
            q_text = process_latex_content(question['question'])
            st.markdown(f"**{i+1}. {q_text}** {q_icon}")
            
            # Already submitted - show results
            if quiz.get('submitted'):
                user_ans = quiz['user_answers'].get(str(q_id))
                correct_ans = question.get('correct')
                
                # Flexible answer checking
                is_correct = check_answer(user_ans, correct_ans, q_type, question)
                
                if q_type == 'multiple_choice':
                    options = question.get('options', [])
                    for j, opt in enumerate(options):
                        opt_text = process_latex_content(opt)
                        if j == correct_ans:
                            st.markdown(f"✅ **{opt_text}** ← Correct")
                        elif j == user_ans:
                            st.markdown(f"❌ ~~{opt_text}~~ ← Your answer")
                        else:
                            st.markdown(f"○ {opt_text}")
                elif q_type == 'true_false':
                    if is_correct:
                        st.success(f"✅ Your answer: {'True' if user_ans else 'False'} - Correct!")
                    else:
                        st.error(f"❌ Your answer: {'True' if user_ans else 'False'} - Expected: {'True' if correct_ans else 'False'}")
                else:
                    # Generic display for other types
                    if is_correct:
                        st.success(f"✅ Your answer: {user_ans} - Correct!")
                    else:
                        st.error(f"❌ Your answer: {user_ans} - Expected: {correct_ans}")
                
                # Show explanation with LaTeX support
                if question.get('explanation'):
                    exp_text = process_latex_content(question['explanation'])
                    st.info(f"💡 {exp_text}")
            
            else:
                # Interactive input based on question type
                render_question_input(quiz, question, q_key, q_type)
            
            st.markdown("---")
        
        # Submit button (only if not submitted)
        if not quiz.get('submitted'):
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                if st.button("📊 Submit Quiz", key=f"{quiz_id}_submit", type="primary", use_container_width=True):
                    # Calculate score using flexible check_answer
                    correct = 0
                    total = len(quiz['questions'])
                    
                    for question in quiz['questions']:
                        q_id = str(question['id'])
                        user_ans = quiz['user_answers'].get(q_id)
                        correct_ans = question.get('correct')
                        q_type = question.get('type', 'short_answer')
                        
                        if check_answer(user_ans, correct_ans, q_type, question):
                            correct += 1
                    
                    quiz['submitted'] = True
                    quiz['score'] = {'correct': correct, 'total': total}
        
        else:
            # Show score
            score = quiz.get('score', {})
            correct = score.get('correct', 0)
            total = score.get('total', 0)
            percentage = (correct / total * 100) if total > 0 else 0
            
            if percentage >= 80:
                st.balloons()
                st.success(f"🎉 Excellent! Score: **{correct}/{total}** ({percentage:.0f}%)")
            elif percentage >= 60:
                st.success(f"👍 Good job! Score: **{correct}/{total}** ({percentage:.0f}%)")
            elif percentage >= 40:
                st.warning(f"📚 Keep practicing! Score: **{correct}/{total}** ({percentage:.0f}%)")
            else:
                st.error(f"💪 Don't give up! Score: **{correct}/{total}** ({percentage:.0f}%)")
            
            # Retry button
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                if st.button("🔄 Retry Quiz", key=f"{quiz_id}_retry", use_container_width=True):
                    quiz['submitted'] = False
                    quiz['user_answers'] = {}
                    quiz['score'] = None


# =============================================================================
# 3-TIER MEMORY TOOL FUNCTIONS (For Gemini)
# =============================================================================
# These are wrapper functions that call the actual implementations in tools.py
# Gemini requires plain Python functions, not LangChain @tool decorated functions

import tools as memory_tools

def save_short_term_memory(content: str, context: str, importance: str = "normal") -> str:
    """
    Save temporary context to short-term memory for THIS SESSION ONLY.
    This memory is specific to the current conversation and won't carry over to new chats.
    
    Args:
        content: The information to remember (what to save)
        context: Why this is relevant (context/reason)
        importance: Priority level - "low", "normal", or "high"
    
    Returns:
        Confirmation message
    """
    return memory_tools.save_short_term_memory.invoke({
        "content": content,
        "context": context,
        "importance": importance
    })


def save_long_term_memory(fact: str, category: str) -> str:
    """
    Save permanent personal information to long-term memory.
    This information NEVER expires unless the user explicitly asks to change it.
    
    Args:
        fact: The personal fact/preference to permanently remember
        category: Type of info - "identity", "preference", "relationship", "date", "skill", "other"
    
    Returns:
        Confirmation that the information is permanently stored
    """
    return memory_tools.save_long_term_memory.invoke({
        "fact": fact,
        "category": category
    })


def save_project_memory(project_name: str, content: str, content_type: str = "note") -> str:
    """
    Save project-related information to project memory.
    
    Args:
        project_name: Name of the project (e.g., "Intro to CS Project")
        content: Information to save (idea, progress, note, requirement, etc.)
        content_type: Type - "goal", "idea", "progress", "note", "requirement", "blocker", "resource"
    
    Returns:
        Confirmation with project tracking info
    """
    return memory_tools.save_project_memory.invoke({
        "project_name": project_name,
        "content": content,
        "content_type": content_type
    })


def query_short_term(query: str) -> str:
    """
    Search short-term memory for context from THIS SESSION ONLY.
    
    Args:
        query: What to search for in recent memory
    
    Returns:
        Relevant short-term memories from current session
    """
    return memory_tools.query_short_term.invoke({"query": query})


def query_long_term(query: str) -> str:
    """
    Search long-term memory for permanent personal facts.
    
    Args:
        query: What personal information to search for
    
    Returns:
        Relevant permanent facts about the user
    """
    return memory_tools.query_long_term.invoke({"query": query})


def query_project(project_name: str = "", query: str = "") -> str:
    """
    Search project memory. Can filter by project name or search across all projects.
    
    Args:
        project_name: Optional - specific project to search (leave empty for all projects)
        query: What to search for within projects
    
    Returns:
        Relevant project information and progress
    """
    return memory_tools.query_project.invoke({
        "project_name": project_name,
        "query": query
    })


def list_all_projects() -> str:
    """
    List all active projects stored in memory.
    
    Returns:
        List of project names with their status and item counts
    """
    return memory_tools.list_all_projects.invoke({})


def update_long_term_memory(old_fact: str, new_fact: str, category: str) -> str:
    """
    Update or correct a long-term memory fact.
    Use when user wants to change previously stored personal information.
    
    Args:
        old_fact: The fact to find and update (approximate match is OK)
        new_fact: The corrected/updated information
        category: Category of the fact being updated
    
    Returns:
        Confirmation of the update
    """
    return memory_tools.update_long_term_memory.invoke({
        "old_fact": old_fact,
        "new_fact": new_fact,
        "category": category
    })


# =============================================================================
# SMART QUERY EXTRACTION
# =============================================================================

def extract_gmail_query(user_message: str) -> str:
    """
    Extract optimal Gmail search query from natural language.
    Uses fast pattern matching instead of AI for speed.
    
    Args:
        user_message: The user's natural language request
    
    Returns:
        Optimized Gmail search query string
    """
    msg_lower = user_message.lower()
    query_parts = []
    
    # Time-based patterns
    if any(w in msg_lower for w in ['today', 'hôm nay']):
        query_parts.append(f"after:{datetime.now().strftime('%Y/%m/%d')}")
    elif any(w in msg_lower for w in ['yesterday', 'hôm qua']):
        yesterday = datetime.now() - timedelta(days=1)
        query_parts.append(f"after:{yesterday.strftime('%Y/%m/%d')}")
    elif any(w in msg_lower for w in ['this week', 'tuần này', 'recent', 'gần đây', 'last few days']):
        query_parts.append("newer_than:7d")
    elif any(w in msg_lower for w in ['this month', 'tháng này']):
        query_parts.append("newer_than:30d")
    # NEW: Handle "new", "latest", "any new" - default to recent 3 days
    elif any(w in msg_lower for w in ['new', 'latest', 'newest', 'mới', 'mới nhất']):
        query_parts.append("newer_than:3d")
    
    # Content type patterns
    if any(w in msg_lower for w in ['unread', 'chưa đọc']):
        query_parts.append("is:unread")
    if any(w in msg_lower for w in ['attachment', 'file', 'đính kèm', 'tệp']):
        query_parts.append("has:attachment")
    if any(w in msg_lower for w in ['starred', 'important', 'quan trọng', 'urgent']):
        query_parts.append("is:starred")
    
    # Meeting/appointment patterns
    if any(w in msg_lower for w in ['meeting', 'appointment', 'schedule', 'calendar', 'invite', 'cuộc họp', 'lịch']):
        query_parts.append("subject:(meeting OR appointment OR calendar OR invite OR schedule)")
    
    # File type patterns
    for ext in ['pdf', 'doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx']:
        if ext in msg_lower:
            query_parts.append(f"filename:{ext}")
            break
    
    # Extract sender patterns
    from_match = re.search(r'from\s+(\S+@\S+|\w+)', msg_lower)
    if from_match:
        query_parts.append(f"from:{from_match.group(1)}")
    
    # Extract subject patterns
    subject_match = re.search(r'(?:subject|about|regarding|re:?)\s+["\']?([^"\']+)["\']?', msg_lower)
    if subject_match:
        query_parts.append(f"subject:{subject_match.group(1).strip()}")
    
    # If we have specific query parts, use them
    if query_parts:
        return ' '.join(query_parts)
    
    # Fallback: If asking about emails in general, default to recent 3 days
    if any(w in msg_lower for w in ['mail', 'email', 'inbox', 'message']):
        return "newer_than:3d"
    
    # Final fallback: extract key terms from user message
    return _simple_query_extraction(user_message)


def _simple_query_extraction(user_message: str) -> str:
    """
    Simple fallback query extraction without AI.
    Removes common filler words and returns key terms.
    """
    # Common words to remove
    stop_words = {
        'find', 'search', 'look', 'for', 'my', 'me', 'i', 'the', 'a', 'an',
        'any', 'some', 'all', 'about', 'from', 'in', 'on', 'at', 'to', 'with',
        'can', 'you', 'please', 'help', 'show', 'get', 'check', 'email', 'emails',
        'gmail', 'mail', 'message', 'messages', 'what', 'where', 'when', 'how',
        'is', 'are', 'was', 'were', 'have', 'has', 'do', 'does', 'did', 'will',
        'would', 'could', 'should', 'may', 'might', 'must', 'shall', 'there',
        'this', 'that', 'these', 'those', 'be', 'been', 'being', 'and', 'or',
        'but', 'if', 'then', 'so', 'because', 'as', 'of', 'by', 'up', 'down'
    }
    
    # Extract words and filter
    words = re.findall(r'\b\w+\b', user_message.lower())
    key_words = [w for w in words if w not in stop_words and len(w) > 2]
    
    # Return top 3-5 keywords
    return ' '.join(key_words[:5]) if key_words else user_message


# =============================================================================
# WEB & YOUTUBE TOOLS
# =============================================================================

def read_web_page(url: str, extract_type: str = "content") -> str:
    """
    Fetch and return content from a web page. Can extract general content or structured data like news headlines.
    
    Args:
        url: The URL of the web page to read.
        extract_type: Type of extraction - "content" for general reading, "news" for headlines/links, "links" for all links.
    
    Returns:
        Extracted content based on extract_type.
    """
    try:
        # Validate URL
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        # For news/links extraction, use BeautifulSoup directly
        if extract_type in ["news", "links"]:
            return _extract_structured_content(url, extract_type)
        
        # Method 1: Try Jina Reader API first (best for most sites)
        try:
            jina_url = f"https://r.jina.ai/{url}"
            response = requests.get(jina_url, timeout=10, headers={
                'User-Agent': 'Mozilla/5.0 (compatible; Googlebot/2.1; +http://www.google.com/bot.html)'
            })
            
            if response.status_code == 200:
                content = response.text.strip()
                
                # Check if we got meaningful content (not empty or error page)
                if content and len(content) > 100 and not content.startswith('Error'):
                    # Limit content size but be more generous
                    if len(content) > 20000:
                        content = content[:20000] + "\n\n[...Content Truncated - Too Long...]"
                    
                    return f"""=== WEB CONTENT ===
Source: {url}
---
{content}
=== END CONTENT ==="""
        except Exception:
            pass  # Fall through to beautifulsoup method
        
        # Method 2: Fallback to beautifulsoup4 (for sites that block Jina)
        try:
            from bs4 import BeautifulSoup
            
            response = requests.get(url, timeout=10, headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            })
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "meta", "link"]):
                script.decompose()
            
            # Try to find main content
            main_content = soup.find('main') or soup.find('article') or soup.find('body')
            if not main_content:
                main_content = soup
            
            # Extract text from paragraphs
            text_parts = []
            
            # Get title
            title = soup.find('h1')
            if title:
                text_parts.append(f"# {title.get_text(strip=True)}\n")
            
            # Get headings and paragraphs
            for elem in main_content.find_all(['h2', 'h3', 'h4', 'p', 'li']):
                text = elem.get_text(strip=True)
                if text and len(text) > 10:  # Only meaningful text
                    if elem.name.startswith('h'):
                        level = int(elem.name[1]) + 1
                        text_parts.append(f"{'#' * level} {text}")
                    else:
                        text_parts.append(text)
            
            content = "\n\n".join(text_parts)
            
            if content and len(content) > 100:
                # Limit content
                if len(content) > 20000:
                    content = content[:20000] + "\n\n[...Content Truncated - Too Long...]"
                
                return f"""=== WEB CONTENT ===
Source: {url}
Method: Direct Extraction
---
{content}
=== END CONTENT ==="""
        
        except ImportError:
            return "Error: beautifulsoup4 not installed. Run: pip install beautifulsoup4"
        except Exception as e:
            pass  # Fall through to error
        
        # If both methods failed
        return f"Error: Could not retrieve content from {url}. The page may be blocked or unavailable."
    
    except Exception as e:
        return f"Error: Failed to read web page. {str(e)[:80]}"


def _extract_structured_content(url: str, extract_type: str) -> str:
    """
    Extract structured content like news headlines or links from a webpage.
    
    Args:
        url: The URL to extract from
        extract_type: "news" for headlines, "links" for all links
    
    Returns:
        Formatted list of extracted items
    """
    try:
        from bs4 import BeautifulSoup
        from urllib.parse import urljoin
        
        response = requests.get(url, timeout=15, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5,vi;q=0.3',
        })
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove unwanted elements
        for elem in soup(['script', 'style', 'nav', 'footer', 'header', 'aside', 'form', 'iframe']):
            elem.decompose()
        
        items = []
        seen_titles = set()
        
        if extract_type == "news":
            # Find all links with meaningful text (likely headlines)
            for a_tag in soup.find_all('a', href=True):
                try:
                    title = a_tag.get_text(strip=True)
                    href = a_tag.get('href', '')
                    
                    # Ensure title and href are strings
                    if not isinstance(title, str):
                        title = str(title)
                    if not isinstance(href, str):
                        href = str(href)
                    
                    # Skip empty values
                    if not title or not href:
                        continue
                    
                    # Filter criteria for news headlines
                    title_len = len(title)
                    if (title_len > 15  # Headlines are usually longer
                        and title_len < 300  # But not too long
                        and title.lower() not in seen_titles
                        and not title.lower().startswith(('đăng nhập', 'login', 'sign', 'menu', 'home', 'trang chủ', 'category', 'tag'))
                        and not href.startswith(('#', 'javascript:', 'mailto:'))
                        and not any(skip in href.lower() for skip in ['/tag/', '/category/', '/author/', '/login', '/register', '/cart'])):
                        
                        # Make URL absolute
                        try:
                            full_url = urljoin(url, href)
                        except Exception:
                            continue
                        
                        seen_titles.add(title.lower())
                        items.append({
                            'title': title,
                            'url': full_url
                        })
                
                except Exception:
                    continue  # Skip problematic links
            
            # Sort by title length (longer titles are usually main headlines)
            try:
                items.sort(key=lambda x: len(x['title']), reverse=True)
            except Exception:
                pass  # If sort fails, keep original order
            
            # Take top items
            items = items[:15]
            
            if items:
                result = f"=== NEWS HEADLINES from {url} ===\n\n"
                for i, item in enumerate(items, 1):
                    try:
                        title = str(item['title'])
                        url_str = str(item['url'])
                        result += f"{i}. **{title}**\n   🔗 {url_str}\n\n"
                    except Exception:
                        continue
                return result if result != f"=== NEWS HEADLINES from {url} ===\n\n" else f"Could not extract news headlines from {url}. The page structure may not be supported."
            else:
                return f"Could not extract news headlines from {url}. The page structure may not be supported."
        
        elif extract_type == "links":
            # Extract all meaningful links
            for a_tag in soup.find_all('a', href=True):
                try:
                    title = a_tag.get_text(strip=True)
                    href = a_tag.get('href', '')
                    
                    # Ensure strings
                    if not isinstance(title, str):
                        title = str(title)
                    if not isinstance(href, str):
                        href = str(href)
                    
                    if title and len(title) > 5 and href and not href.startswith(('#', 'javascript:')):
                        full_url = urljoin(url, href)
                        if title.lower() not in seen_titles:
                            seen_titles.add(title.lower())
                            items.append({'title': title, 'url': full_url})
                except Exception:
                    continue
            
            items = items[:30]
            
            if items:
                result = f"=== LINKS from {url} ===\n\n"
                for i, item in enumerate(items, 1):
                    try:
                        result += f"{i}. {str(item['title'])}: {str(item['url'])}\n"
                    except Exception:
                        continue
                return result
        
        return f"No structured content found at {url}"
        
    except ImportError:
        return "Error: beautifulsoup4 not installed. Run: pip install beautifulsoup4"
    except Exception as e:
        return f"Error extracting content from {url}: {str(e)[:80]}"


def get_news_headlines(url: str, count: int = 10) -> str:
    """
    Extract news headlines from a news website. Perfect for sites like VnExpress, BBC, CNN, etc.
    
    Args:
        url: The URL of the news website (e.g., "vnexpress.net", "bbc.com/news")
        count: Number of headlines to return (default 10, max 20)
    
    Returns:
        List of news headlines with their links
    """
    # Validate and normalize URL
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    
    # Limit count
    count = min(max(count, 1), 20)
    
    result = _extract_structured_content(url, "news")
    
    # If we got results, potentially trim to requested count
    if "===" in result and "Could not" not in result:
        lines = result.split('\n')
        header = lines[0:2]
        items = []
        current_item = []
        
        for line in lines[2:]:
            if line.strip().startswith(tuple(f"{i}." for i in range(1, 21))):
                if current_item:
                    items.append('\n'.join(current_item))
                current_item = [line]
            elif current_item:
                current_item.append(line)
        
        if current_item:
            items.append('\n'.join(current_item))
        
        # Return only requested count
        items = items[:count]
        
        return '\n'.join(header) + '\n' + '\n'.join(items)
    
    return result


def watch_youtube(video_url: str) -> str:
    """
    Fetch and return the transcript of a YouTube video with 60-second time blocks.
    Uses the youtube-transcript-api library (v0.6+) with proper language support.
    
    Args:
        video_url: The URL of the YouTube video (supports youtube.com/watch?v= and youtu.be/ formats).
    
    Returns:
        Formatted transcript with timestamps at 60-second intervals.
    """
    try:
        # Extract video ID from URL
        video_id = None
        
        # Pattern for youtube.com/watch?v=VIDEO_ID
        match = re.search(r'[?&]v=([a-zA-Z0-9_-]{11})', video_url)
        if match:
            video_id = match.group(1)
        
        # Pattern for youtu.be/VIDEO_ID
        if not video_id:
            match = re.search(r'youtu\.be/([a-zA-Z0-9_-]{11})', video_url)
            if match:
                video_id = match.group(1)
        
        if not video_id:
            return "Error: Could not extract video ID from URL. Please use a valid YouTube URL."
        
        # Initialize the API client (v0.6+)
        ytt = YouTubeTranscriptApi()
        transcript = None
        language_used = None
        
        # Try to fetch transcript in order: English, Vietnamese, then any available
        for lang_code, lang_name in [('en', 'English'), ('vi', 'Vietnamese')]:
            try:
                transcript = ytt.fetch(video_id, languages=[lang_code])
                language_used = lang_name
                break
            except Exception:
                continue
        
        # If specific languages failed, try to get any available transcript
        if not transcript:
            try:
                available = ytt.list(video_id)
                # Use find_transcript with fallback - iterate over available
                for t in available:
                    try:
                        transcript = t.fetch()
                        language_used = t.language
                        break
                    except Exception:
                        continue
            except Exception as e:
                return f"Error: Could not retrieve transcript for this video. {str(e)[:50]}"
        
        if not transcript or len(transcript) == 0:
            return "Error: Transcript was empty or could not be processed."
        
        # Group transcript into 60-second chunks
        formatted_parts = []
        current_minute = 0
        current_texts = []
        
        for entry in transcript:
            # Handle both dict and object formats
            if hasattr(entry, 'text'):
                text = entry.text
                start_time = entry.start
            else:
                text = entry.get('text', '')
                start_time = entry.get('start', 0)
            
            if not text:
                continue
            
            entry_minute = int(start_time // 60)
            
            # If we've moved to a new minute block
            if entry_minute > current_minute:
                # Save the previous block if it has content
                if current_texts:
                    timestamp = f"[{current_minute:02d}:00]"
                    formatted_parts.append(f"{timestamp} {' '.join(current_texts)}")
                
                # Start new block
                current_minute = entry_minute
                current_texts = [text]
            else:
                current_texts.append(text)
        
        # Don't forget the last block
        if current_texts:
            timestamp = f"[{current_minute:02d}:00]"
            formatted_parts.append(f"{timestamp} {' '.join(current_texts)}")
        
        if not formatted_parts:
            return "Error: Transcript was empty or could not be processed."
        
        formatted_transcript = "\n\n".join(formatted_parts)
        
        # Return formatted output with language info
        return f"""=== YOUTUBE TRANSCRIPT ===
Video ID: {video_id}
Language: {language_used or 'Unknown'}
---
{formatted_transcript}
=== END TRANSCRIPT ==="""
    
    except ImportError:
        return "Error: youtube-transcript-api is not installed. Run: pip install youtube-transcript-api"
    except Exception as e:
        return f"Error: Could not retrieve transcript for this video. {str(e)[:100]}"


def recall_memory(doc_id: str) -> str:
    """
    Retrieve the full content of a specific memory by its document ID.
    Use this when you need complete details from a memory summary shown in context.
    
    Args:
        doc_id: The document ID from the memory context (e.g., "email_20241203_abc123")
    
    Returns:
        Full text content of the memory, or error message if not found.
    """
    memory = get_memory_by_id(doc_id)
    
    if not memory:
        return f"Memory not found with ID: {doc_id}"
    
    metadata = memory.get('metadata', {})
    text = memory.get('text', '')
    
    # Format output with metadata header
    output_parts = ["=== FULL MEMORY CONTENT ==="]
    
    if metadata.get('subject'):
        output_parts.append(f"Subject: {metadata['subject']}")
    if metadata.get('source'):
        output_parts.append(f"Source: {metadata['source']}")
    if metadata.get('sender'):
        output_parts.append(f"From: {metadata['sender']}")
    if metadata.get('date'):
        output_parts.append(f"Date: {metadata['date']}")
    
    output_parts.append("---")
    output_parts.append(text)
    output_parts.append("=== END MEMORY ===")
    
    return "\n".join(output_parts)


def search_memory(query: str) -> str:
    """
    Search through stored memories for relevant information.
    Returns summaries of matching memories. Use recall_memory to get full content.
    
    Args:
        query: Search query describing what you're looking for
    
    Returns:
        List of relevant memory summaries with their IDs
    """
    memories = query_memory(query, n_results=5, threshold=1.0)  # More lenient for explicit search
    
    if not memories:
        return "No memories found matching your query."
    
    output_parts = [f"Found {len(memories)} relevant memories:\n"]
    
    for i, memory in enumerate(memories, 1):
        metadata = memory.get('metadata', {})
        relevance = memory.get('relevance', 0)
        doc_id = memory.get('doc_id', 'unknown')
        summary = memory.get('summary', 'No summary')
        
        output_parts.append(f"{i}. [{relevance}% match] {summary}")
        output_parts.append(f"   Source: {metadata.get('source', 'Unknown')} | ID: {doc_id}")
        output_parts.append("")
    
    output_parts.append("Use recall_memory(doc_id) to get full content of any memory.")
    
    return "\n".join(output_parts)


# =============================================================================
# TASK 2: IN-MEMORY PARSERS
# =============================================================================

def parse_attachment(file_data: bytes, filename: str) -> Optional[str]:
    """
    Parse attachment content without saving to disk.
    
    Args:
        file_data: Raw bytes of the attachment
        filename: Original filename (used to determine file type)
    
    Returns:
        Extracted text content as string, or None if parsing fails
    """
    filename_lower = filename.lower()
    
    try:
        # Handle Excel files (.xlsx, .xls)
        if filename_lower.endswith(('.xlsx', '.xls')):
            return parse_excel(file_data, filename)
        
        # Handle Word documents (.docx)
        elif filename_lower.endswith('.docx'):
            return parse_docx(file_data)
        
        # Handle PDF files (.pdf)
        elif filename_lower.endswith('.pdf'):
            return parse_pdf(file_data)
        
        # Handle plain text files
        elif filename_lower.endswith(('.txt', '.csv', '.md')):
            return file_data.decode('utf-8', errors='ignore')
        
        else:
            return f"[Unsupported file type: {filename}]"
    
    except Exception as e:
        return f"[Error parsing {filename}: {str(e)}]"


def parse_excel(file_data: bytes, filename: str) -> str:
    """
    Parse Excel file from bytes and return CSV string summary.
    """
    # Use BytesIO to read Excel without saving to disk
    buffer = io.BytesIO(file_data)
    
    # Read all sheets
    excel_file = pd.ExcelFile(buffer)
    sheet_names = excel_file.sheet_names
    
    result_parts = [f"📊 Excel File: {filename}"]
    result_parts.append(f"Sheets: {', '.join(sheet_names)}\n")
    
    for sheet_name in sheet_names:
        df = pd.read_excel(buffer, sheet_name=sheet_name)
        result_parts.append(f"--- Sheet: {sheet_name} ---")
        result_parts.append(f"Shape: {df.shape[0]} rows × {df.shape[1]} columns")
        result_parts.append(f"Columns: {', '.join(df.columns.astype(str))}")
        
        # Convert to CSV string (limit rows for summary)
        if len(df) > 50:
            result_parts.append("(Showing first 50 rows)")
            csv_str = df.head(50).to_csv(index=False)
        else:
            csv_str = df.to_csv(index=False)
        
        result_parts.append(csv_str)
        result_parts.append("")
    
    return "\n".join(result_parts)


def parse_docx(file_data: bytes) -> str:
    """
    Parse Word document from bytes and extract all text.
    """
    buffer = io.BytesIO(file_data)
    doc = Document(buffer)
    
    # Extract text from all paragraphs
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    
    # Also extract text from tables
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                table_texts.append(row_text)
    
    result_parts = paragraphs
    if table_texts:
        result_parts.append("\n--- Tables ---")
        result_parts.extend(table_texts)
    
    return "\n".join(result_parts)


def parse_pdf(file_data: bytes) -> str:
    """
    Parse PDF file from bytes and extract text.
    Uses pypdf if available.
    """
    try:
        from pypdf import PdfReader
        
        buffer = io.BytesIO(file_data)
        reader = PdfReader(buffer)
        
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i + 1} ---")
                text_parts.append(page_text)
        
        return "\n".join(text_parts) if text_parts else "[No text extracted from PDF]"
    
    except ImportError:
        return "[PDF parsing skipped - pypdf not installed]"
    except Exception as e:
        return f"[Error parsing PDF: {str(e)}]"


# =============================================================================
# TASK 3: MEMORY LAYER (ChromaDB) - Thread-Safe Version
# =============================================================================

import threading
import time

# ChromaDB persistent storage path
CHROMA_PATH = "./emo_memory"

# Lock for thread-safe database access
_chroma_lock = threading.Lock()
_chroma_client = None
_chroma_collection = None
MAX_RETRIES = 3
RETRY_DELAY = 0.5  # seconds


def get_chroma_client():
    """
    Get or create the ChromaDB client with retry logic.
    Thread-safe implementation to prevent "readonly database" errors.
    """
    global _chroma_client
    
    with _chroma_lock:
        if _chroma_client is None:
            # Create client with sensible timeouts and settings
            try:
                import chromadb
                
                # Ensure directory exists with proper permissions
                os.makedirs(CHROMA_PATH, exist_ok=True)
                
                _chroma_client = chromadb.PersistentClient(path=CHROMA_PATH)
            except Exception as e:
                st.error(f"Failed to initialize ChromaDB: {str(e)}")
                _chroma_client = None
        
        return _chroma_client


def get_chroma_collection():
    """
    Get or create the ChromaDB collection for Emo's memory.
    Uses Streamlit session state to cache the collection.
    Includes retry logic to handle database locks.
    """
    if 'chroma_collection' not in st.session_state:
        collection = None
        
        for attempt in range(MAX_RETRIES):
            try:
                client = get_chroma_client()
                if client is None:
                    raise Exception("ChromaDB client not initialized")
                
                with _chroma_lock:
                    collection = client.get_or_create_collection(
                        name="emo_memory",
                        metadata={"description": "Emo's long-term memory from emails and attachments"}
                    )
                    break  # Success
            
            except Exception as e:
                if attempt < MAX_RETRIES - 1:
                    time.sleep(RETRY_DELAY * (attempt + 1))  # Exponential backoff
                else:
                    # Last attempt failed, log and raise
                    st.warning(f"ChromaDB collection issue (attempt {attempt + 1}/{MAX_RETRIES}): {str(e)}")
                    return None
        
        st.session_state.chroma_collection = collection
    
    return st.session_state.chroma_collection


def generate_summary(text: str, max_length: int = 150) -> str:
    """
    Generate a concise summary of text for memory storage.
    Uses fast heuristic extraction instead of AI for speed.
    
    Args:
        text: The text to summarize
        max_length: Maximum length of summary (default: 150 chars)
    
    Returns:
        A concise summary string
    """
    # If text is already short, use it directly
    if len(text) <= max_length:
        return text.strip()
    
    # Fast extraction: get first meaningful content
    # Remove excessive whitespace
    clean_text = ' '.join(text.split())
    
    # For emails, try to get the first substantive paragraph
    lines = clean_text.split('\n')
    content_lines = []
    for line in lines:
        line = line.strip()
        # Skip headers and empty lines
        if line and not line.startswith(('Subject:', 'From:', 'Date:', 'To:', 'CC:')):
            content_lines.append(line)
        if len(' '.join(content_lines)) > max_length:
            break
    
    if content_lines:
        summary = ' '.join(content_lines)
    else:
        summary = clean_text
    
    # Truncate at sentence boundary if possible
    if len(summary) > max_length:
        truncated = summary[:max_length]
        last_period = truncated.rfind('.')
        last_question = truncated.rfind('?')
        last_exclaim = truncated.rfind('!')
        last_sentence = max(last_period, last_question, last_exclaim)
        
        if last_sentence > max_length // 2:
            return truncated[:last_sentence + 1]
        return truncated[:max_length-3] + "..."
    
    return summary


def save_to_memory(text: str, metadata: dict) -> str:
    """
    Save a document to ChromaDB memory with auto-generated summary.
    Includes retry logic and error handling for database access.
    
    Args:
        text: The text content to store
        metadata: Dictionary with metadata (source, subject, date, etc.)
    
    Returns:
        The document ID that was created, or error string if failed
    """
    try:
        collection = get_chroma_collection()
        
        if collection is None:
            return "[Error: Memory database not available]"
        
        # Generate a unique ID based on content hash and timestamp
        import hashlib
        from datetime import datetime
        
        content_hash = hashlib.md5(text.encode()).hexdigest()[:8]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_id = f"{metadata.get('type', 'doc')}_{timestamp}_{content_hash}"
        
        # Generate summary for efficient retrieval
        summary = generate_summary(text)
        
        # Ensure metadata values are valid types for ChromaDB
        clean_metadata = {}
        for key, value in metadata.items():
            if isinstance(value, (str, int, float, bool)):
                clean_metadata[key] = value
            else:
                clean_metadata[key] = str(value)
        
        # Store summary in metadata for quick access
        clean_metadata['summary'] = summary
        clean_metadata['full_text_length'] = len(text)
        
        # Upsert with retry logic
        for attempt in range(MAX_RETRIES):
            try:
                with _chroma_lock:
                    collection.upsert(
                        ids=[doc_id],
                        documents=[text],
                        metadatas=[clean_metadata]
                    )
                return doc_id
            
            except Exception as e:
                if attempt < MAX_RETRIES - 1:
                    time.sleep(RETRY_DELAY * (attempt + 1))
                else:
                    return f"[Error saving to memory: {str(e)[:50]}]"
    
    except Exception as e:
        return f"[Error in save_to_memory: {str(e)[:50]}]"


# Distance threshold for relevance filtering (lower = more similar)
# ChromaDB uses L2 distance by default: 0 = identical, >2.0 = likely irrelevant
# Adjusted to 1.0 to allow reasonable matches while filtering noise
MEMORY_DISTANCE_THRESHOLD = 1.0


def query_memory(query_text: str, n_results: int = 3, threshold: float = MEMORY_DISTANCE_THRESHOLD) -> list[dict]:
    """
    Query ChromaDB for the most relevant memories.
    Only returns results that pass the relevance threshold.
    Includes exception handling for database access issues.
    
    Args:
        query_text: The search query
        n_results: Maximum number of results to return (default: 3)
        threshold: Maximum distance for relevance (default: 0.8)
    
    Returns:
        List of dictionaries with 'summary', 'metadata', 'distance', and 'doc_id' keys
    """
    try:
        collection = get_chroma_collection()
        
        if collection is None:
            return []
        
        # Check if collection has any documents
        if collection.count() == 0:
            return []
        
        # Query for similar documents with retry logic
        for attempt in range(MAX_RETRIES):
            try:
                with _chroma_lock:
                    results = collection.query(
                        query_texts=[query_text],
                        n_results=min(n_results, collection.count()),
                        include=['documents', 'metadatas', 'distances']
                    )
                    break
            except Exception as e:
                if attempt < MAX_RETRIES - 1:
                    time.sleep(RETRY_DELAY * (attempt + 1))
                else:
                    return []
        
        # Format results with relevance filtering
        memories = []
        if results['documents'] and results['documents'][0]:
            for i, doc in enumerate(results['documents'][0]):
                distance = results['distances'][0][i] if results['distances'] else 1.0
                
                # Skip if below relevance threshold
                if distance > threshold:
                    continue
                
                metadata = results['metadatas'][0][i] if results['metadatas'] else {}
                doc_id = results['ids'][0][i] if results['ids'] else None
                
                # Use summary if available, otherwise truncate
                summary = metadata.get('summary', doc[:150] + '...' if len(doc) > 150 else doc)
                
                memory = {
                    'summary': summary,
                    'metadata': metadata,
                    'distance': distance,
                    'doc_id': doc_id,
                    'relevance': round((1 - distance) * 100)  # Convert to percentage
                }
                memories.append(memory)
        
        return memories
    
    except Exception as e:
        return []


def get_memory_by_id(doc_id: str) -> dict | None:
    """
    Retrieve full memory content by document ID.
    Includes exception handling for database access issues.
    
    Args:
        doc_id: The document ID to retrieve
    
    Returns:
        Dictionary with full 'text' and 'metadata', or None if not found
    """
    try:
        collection = get_chroma_collection()
        
        if collection is None:
            return None
        
        for attempt in range(MAX_RETRIES):
            try:
                with _chroma_lock:
                    result = collection.get(
                        ids=[doc_id],
                        include=['documents', 'metadatas']
                    )
                    
                    if result['documents'] and result['documents'][0]:
                        return {
                            'text': result['documents'][0],
                            'metadata': result['metadatas'][0] if result['metadatas'] else {}
                        }
                    return None
            
            except Exception as e:
                if attempt < MAX_RETRIES - 1:
                    time.sleep(RETRY_DELAY * (attempt + 1))
                else:
                    return None
    
    except Exception:
        return None


def format_memories_for_context(memories: list[dict]) -> str:
    """
    Format retrieved memories into a compact context string for the LLM.
    
    IMPORTANT: Skips email memories because they only contain summaries (150 chars).
    For emails, the AI MUST call Gmail tools to get full content.
    """
    if not memories:
        return ""
    
    # Compact format optimized for small context models
    context_parts = ["[Memory Context]"]
    has_attachments = False
    has_emails = False
    included_count = 0
    
    for i, memory in enumerate(memories, 1):
        metadata = memory.get('metadata', {})
        source = metadata.get('source', '')
        subject = metadata.get('subject', '')
        doc_type = metadata.get('type', '')
        filename = metadata.get('filename', '')
        relevance = memory.get('relevance', 0)
        doc_id = memory.get('doc_id', '')
        
        # SKIP email memories - they only have summaries, not full content!
        # The AI must fetch fresh from Gmail to avoid hallucination
        if doc_type == 'email' or source == 'Gmail':
            has_emails = True
            continue  # Skip - force AI to use Gmail tools
        
        included_count += 1
        
        # Build compact header
        header_parts = [f"#{included_count}"]
        if relevance:
            header_parts.append(f"{relevance}%")
        
        # Mark attachments clearly
        if doc_type == 'attachment' or filename:
            header_parts.append(f"📎 ATTACHMENT: {filename}")
            has_attachments = True
        elif subject:
            header_parts.append(f"'{subject}'")
        
        if source:
            header_parts.append(f"from {source}")
        
        header = " | ".join(header_parts)
        summary = memory.get('summary', 'No summary available')
        
        context_parts.append(f"{header}\n→ {summary}")
        if doc_id:
            context_parts.append(f"  [ID: {doc_id}]")
    
    # Critical warning about emails
    if has_emails:
        context_parts.append("\n⚠️ EMAIL CONTENT NOT IN MEMORY - You MUST call `quick_gmail_search` to get full email content! DO NOT use summaries!")
    
    # Strong reminder for attachments
    if has_attachments:
        context_parts.append("\n⚠️ ATTACHMENTS DETECTED: Use `recall_memory(doc_id)` to read file contents!")
    elif included_count > 0:
        context_parts.append("\n(Use recall_memory(doc_id) for full content)")
    
    if included_count == 0 and has_emails:
        return "[Memory Context]\n⚠️ For email content, you MUST call `quick_gmail_search(query)` - memory only has summaries!"
    
    return "\n".join(context_parts)


# =============================================================================
# TASK 4: AGENT TOOL (Gmail + Memory)
# =============================================================================

def get_email_body(payload: dict) -> str:
    """
    Extract email body from Gmail message payload.
    Handles deeply nested multipart messages (common with attachments).
    
    Strategy: Recursively collect ALL text parts, prefer plain text over HTML.
    """
    plain_texts = []
    html_texts = []
    
    def extract_parts(part):
        """Recursively extract text from all parts."""
        mime_type = part.get('mimeType', '')
        
        # Check if this part has direct body data
        if part.get('body', {}).get('data'):
            try:
                decoded = base64.urlsafe_b64decode(part['body']['data']).decode('utf-8', errors='ignore')
                if decoded.strip():  # Only add non-empty content
                    if mime_type == 'text/plain':
                        plain_texts.append(decoded)
                    elif mime_type == 'text/html':
                        html_texts.append(decoded)
            except Exception:
                pass
        
        # Recursively process nested parts (multipart/alternative, multipart/mixed, etc.)
        if 'parts' in part:
            for subpart in part['parts']:
                extract_parts(subpart)
    
    # Start extraction from the root payload
    extract_parts(payload)
    
    # Prefer plain text if available
    if plain_texts:
        return '\n\n'.join(plain_texts)
    
    # Fall back to HTML, but clean it up
    if html_texts:
        import re
        html_content = '\n\n'.join(html_texts)
        
        # Convert common HTML to text
        # Remove script and style tags completely
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Convert line breaks
        html_content = re.sub(r'<br\s*/?>', '\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</p>', '\n\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</div>', '\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</li>', '\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</tr>', '\n', html_content, flags=re.IGNORECASE)
        
        # Remove all remaining HTML tags
        html_content = re.sub(r'<[^>]+>', '', html_content)
        
        # Clean up HTML entities
        html_content = html_content.replace('&nbsp;', ' ')
        html_content = html_content.replace('&amp;', '&')
        html_content = html_content.replace('&lt;', '<')
        html_content = html_content.replace('&gt;', '>')
        html_content = html_content.replace('&quot;', '"')
        html_content = html_content.replace('&#39;', "'")
        
        # Clean up excessive whitespace while preserving paragraph structure
        lines = html_content.split('\n')
        cleaned_lines = [line.strip() for line in lines]
        html_content = '\n'.join(line for line in cleaned_lines if line)
        
        return html_content
    
    return ""


def get_attachments(service, message_id: str, payload: dict) -> list[tuple[str, bytes]]:
    """
    Extract all attachments from an email.
    
    Returns:
        List of tuples: (filename, file_data_bytes)
    """
    attachments = []
    
    def process_parts(parts):
        for part in parts:
            filename = part.get('filename', '')
            
            if filename:
                # This is an attachment
                attachment_id = part.get('body', {}).get('attachmentId')
                
                if attachment_id:
                    # Fetch attachment data
                    att = service.users().messages().attachments().get(
                        userId='me',
                        messageId=message_id,
                        id=attachment_id
                    ).execute()
                    
                    file_data = base64.urlsafe_b64decode(att['data'])
                    attachments.append((filename, file_data))
                
                elif part.get('body', {}).get('data'):
                    # Small attachment included directly
                    file_data = base64.urlsafe_b64decode(part['body']['data'])
                    attachments.append((filename, file_data))
            
            # Check nested parts
            if 'parts' in part:
                process_parts(part['parts'])
    
    if 'parts' in payload:
        process_parts(payload['parts'])
    
    return attachments


def check_gmail_and_learn(query: str, max_emails: int = 2, skip_attachments: bool = True) -> str:
    """
    Tool function: Search Gmail and return email content.
    NO LONGER SAVES TO MEMORY - fetches fresh each time to prevent hallucination.
    
    Args:
        query: Natural language search query (e.g., "find my upcoming appointments")
        max_emails: Maximum number of emails to fetch (default: 2 for speed)
        skip_attachments: If True, skip attachment parsing for faster results (default: True)
    
    Returns:
        Email content fetched directly from Gmail
    """
    try:
        service = get_gmail_service()
        
        # Fast query extraction - no AI call needed
        optimized_query = extract_gmail_query(query)
        
        # Search Gmail with optimized query
        results = service.users().messages().list(
            userId='me',
            q=optimized_query,
            maxResults=max_emails
        ).execute()
        
        messages = results.get('messages', [])
        
        if not messages:
            return f"No emails found. Search query used: '{optimized_query}' (from: '{query}')"
        
        email_count = 0
        subjects = []
        
        # Fetch message content
        for msg_info in messages:
            # Fetch full message
            message = service.users().messages().get(
                userId='me',
                id=msg_info['id'],
                format='full'
            ).execute()
            
            # Extract headers
            headers = {h['name']: h['value'] for h in message['payload'].get('headers', [])}
            subject = headers.get('Subject', 'No Subject')
            
            subjects.append(subject)
            
            # Extract email body
            body = get_email_body(message['payload'])
            
            if body:
                email_count += 1
                # NOTE: No longer saving to memory - fetch fresh each time
        
        # Format response
        subject_list = ', '.join(f"'{s}'" for s in subjects[:3])
        if len(subjects) > 3:
            subject_list += f" and {len(subjects) - 3} more"
        
        return (
            f"✅ Found {email_count} email(s).\n"
            f"📧 Subjects: {subject_list}\n"
            f"🔍 Search used: `{optimized_query}`\n"
            f"💡 Use `quick_gmail_search` to see full content."
        )
    
    except Exception as e:
        return f"❌ Error accessing Gmail: {str(e)}"


def quick_gmail_search(query: str, max_results: int = 3) -> str:
    """
    Search Gmail and return FULL email content.
    CACHES results so you can refer to them by index later (e.g., "email 1").
    
    Args:
        query: Natural language search query
        max_results: Maximum number of results (default: 3)
    
    Returns:
        Complete email content with full body text
    """
    try:
        service = get_gmail_service()
        optimized_query = extract_gmail_query(query)
        
        results = service.users().messages().list(
            userId='me',
            q=optimized_query,
            maxResults=max_results
        ).execute()
        
        messages = results.get('messages', [])
        
        if not messages:
            return f"No emails found for: '{query}'"
        
        # Initialize cache if needed
        if 'email_cache' not in st.session_state:
            st.session_state.email_cache = {}
            
        email_contents = []
        current_search_results = []
        
        for i, msg_info in enumerate(messages, 1):
            # Fetch FULL message
            message = service.users().messages().get(
                userId='me',
                id=msg_info['id'],
                format='full'
            ).execute()
            
            headers = {h['name']: h['value'] for h in message['payload'].get('headers', [])}
            subject = headers.get('Subject', 'No Subject')
            sender = headers.get('From', 'Unknown')
            date = headers.get('Date', 'Unknown')
            
            # Cache this email's ID and basic info
            current_search_results.append({
                'index': i,
                'id': msg_info['id'],
                'subject': subject,
                'sender': sender
            })
            
            # Get the FULL email body
            body = get_email_body(message['payload'])
            
            # Check for attachments
            attachment_names = []
            def find_attachments(payload):
                if 'parts' in payload:
                    for part in payload['parts']:
                        if part.get('filename'):
                            attachment_names.append(part['filename'])
                        if 'parts' in part:
                            find_attachments(part)
            find_attachments(message['payload'])
            
            # Build email output
            email_output = [
                f"--- EMAIL #{i} ---",
                f"**Subject:** {subject}",
                f"**From:** {sender}",
                f"**Date:** {date}"
            ]
            
            if attachment_names:
                email_output.append(f"**Attachments ({len(attachment_names)}):** {', '.join(attachment_names)}")
            
            email_output.append("")
            email_output.append("**Full Content:**")
            email_output.append(body if body else "(No text content)")
            email_output.append("")
            
            email_contents.append("\n".join(email_output))
        
        # Update session cache with these results
        st.session_state.email_cache = current_search_results
        
        return f"Found {len(messages)} email(s) (Cached indices 1-{len(messages)}):\n\n" + "\n".join(email_contents)
    
    except Exception as e:
        return f"❌ Error: {str(e)}"


def get_email_by_index(index: int) -> str:
    """
    Fetch email content by its index from the LAST search results.
    Use this when user says "email 2", "the first one", "number 3", etc.
    
    Args:
        index: The index number (1-based) from the previous search list
    
    Returns:
        Full content of the specific email
    """
    try:
        if 'email_cache' not in st.session_state or not st.session_state.email_cache:
            return "❌ No recent search results found. Please search for emails first."
        
        # Find the email with the matching index
        target_email = None
        for email in st.session_state.email_cache:
            if email['index'] == index:
                target_email = email
                break
        
        if not target_email:
            return f"❌ Email #{index} not found in recent results (available: 1-{len(st.session_state.email_cache)})"
            
        # Save this as the last viewed email for context
        st.session_state.last_viewed_email_index = index
        
        # Fetch the specific email by ID
        service = get_gmail_service()
        message = service.users().messages().get(
            userId='me',
            id=target_email['id'],
            format='full'
        ).execute()
        
        # Extract content
        headers = {h['name']: h['value'] for h in message['payload'].get('headers', [])}
        subject = headers.get('Subject', 'No Subject')
        sender = headers.get('From', 'Unknown')
        date = headers.get('Date', 'Unknown')
        
        body = get_email_body(message['payload'])
        
        attachment_names = []
        def find_attachments(payload):
            if 'parts' in payload:
                for part in payload['parts']:
                    if part.get('filename'):
                        attachment_names.append(part['filename'])
                    if 'parts' in part:
                        find_attachments(part)
        find_attachments(message['payload'])
        
        email_output = [
            f"**Email #{index}**",
            f"**Subject:** {subject}",
            f"**From:** {sender}",
            f"**Date:** {date}"
        ]
        
        if attachment_names:
            email_output.append(f"**Attachments:** {', '.join(attachment_names)}")
        
        email_output.append("\n---\n")
        
        # Escape the body content to prevent markdown header interpretation
        # Replace lines that could be interpreted as headers
        if body:
            # Wrap each line that could be a header in a way that prevents markdown interpretation
            safe_body = body.replace('\n\n', '\n\n<!-- -->\n\n')  # Break up setext headers
            email_output.append(safe_body)
        else:
            email_output.append("(No text content)")
        
        return "\n".join(email_output)
        
    except Exception as e:
        return f"❌ Error fetching email #{index}: {str(e)}"


def analyze_attachment(email_index: int, attachment_index: int = 1) -> str:
    """
    Analyze attachment(s) from a cached email and return a 'cool' summary.
    Optimized for SLMs - returns a structured 'File DNA' card instead of full text.
    
    Args:
        email_index: The index of the email (1, 2, 3...)
        attachment_index: The index of the attachment (default: 1). Set to 0 for ALL attachments.
    
    Returns:
        A cool, structured summary of the attachment(s).
    """
    try:
        if 'email_cache' not in st.session_state or not st.session_state.email_cache:
            return "❌ No recent search results. Search for emails first!"
        
        # 1. Find the email
        target_email = None
        for email in st.session_state.email_cache:
            if email['index'] == email_index:
                target_email = email
                break
        
        if not target_email:
            return f"❌ Email #{email_index} not found."
            
        # 2. Fetch email to get attachment IDs
        service = get_gmail_service()
        message = service.users().messages().get(
            userId='me',
            id=target_email['id'],
            format='full'
        ).execute()
        
        # 3. Find attachments
        all_attachments = []
        
        def collect_attachments(payload):
            if 'parts' in payload:
                for part in payload['parts']:
                    if part.get('filename'):
                        all_attachments.append(part)
                    if 'parts' in part:
                        collect_attachments(part)
        
        collect_attachments(message['payload'])
        
        if not all_attachments:
            return f"❌ No attachments found in Email #{email_index}."
            
        # Determine which attachments to process
        targets = []
        if attachment_index == 0:
            targets = all_attachments
        elif 1 <= attachment_index <= len(all_attachments):
            targets = [all_attachments[attachment_index - 1]]
        else:
            return f"❌ Attachment #{attachment_index} not found (Email has {len(all_attachments)} attachments)."
            
        results = []
        
        for att in targets:
            filename = att['filename']
            attachment_id = att['body'].get('attachmentId')
            
            # Download content
            if attachment_id:
                att_data = service.users().messages().attachments().get(
                    userId='me',
                    messageId=target_email['id'],
                    id=attachment_id
                ).execute()
                file_data = base64.urlsafe_b64decode(att_data['data'])
            elif att['body'].get('data'):
                file_data = base64.urlsafe_b64decode(att['body']['data'])
            else:
                results.append(f"❌ Could not download '{filename}'")
                continue
                
            # Parse content
            text_content = parse_attachment(file_data, filename)
            
            if not text_content or text_content.startswith("[Error"):
                results.append(f"❌ Could not parse '{filename}'")
                continue
                
            # Generate 'Cool' SLM Summary (File DNA)
            file_size_kb = len(file_data) / 1024
            
            # Determine 'Vibe' (File Type Icon)
            icon = "📄"
            if filename.endswith(('.xls', '.xlsx')): icon = "📊"
            elif filename.endswith('.pdf'): icon = "📕"
            elif filename.endswith(('.jpg', '.png')): icon = "🖼️"
            elif filename.endswith('.zip'): icon = "📦"
            elif filename.endswith('.pptx'): icon = "📽️"
            
            # Extract 'Essence' (First 500 chars clean)
            clean_text = " ".join(text_content.split())[:500] + "..."
            
            # Heuristic Keyword Extraction
            words = [w for w in text_content.split() if len(w) > 5 and w[0].isupper()]
            top_keywords = list(set(words))[:5]
            
            card = f"""
[analyze_attachment result]:
╭──────────────────────────────────────╮
│  {icon}  **FILE DNA: {filename}**
╰──────────────────────────────────────╯
**Size:** {file_size_kb:.1f} KB

**📄 ACTUAL CONTENT FROM FILE (first 500 chars):**
"{clean_text}"

**🔑 Keywords Found:** {', '.join([f"`{k}`" for k in top_keywords]) if top_keywords else "None"}

⚠️ THIS IS THE REAL FILE CONTENT - DO NOT MAKE UP ADDITIONAL DETAILS
"""
            results.append(card)
            
        return "\n".join(results)

    except Exception as e:
        return f"❌ Error analyzing attachment: {str(e)}"


def get_full_email(query: str) -> str:
    """
    Fetch the FULL content of an email matching the query.
    Use this when user wants complete email details, not just a summary.
    
    Args:
        query: Natural language search query to find the specific email
    
    Returns:
        Complete email content with all details (body, attendees, links, etc.)
    """
    try:
        service = get_gmail_service()
        optimized_query = extract_gmail_query(query)
        
        # Search for the email
        results = service.users().messages().list(
            userId='me',
            q=optimized_query,
            maxResults=1  # Get just the most relevant email
        ).execute()
        
        messages = results.get('messages', [])
        
        if not messages:
            return f"No email found matching: '{query}'"
        
        # Fetch the FULL message (not just metadata)
        message = service.users().messages().get(
            userId='me',
            id=messages[0]['id'],
            format='full'
        ).execute()
        
        # Extract headers
        headers = {h['name']: h['value'] for h in message['payload'].get('headers', [])}
        subject = headers.get('Subject', 'No Subject')
        sender = headers.get('From', 'Unknown')
        date = headers.get('Date', 'Unknown')
        to = headers.get('To', '')
        cc = headers.get('Cc', '')
        
        # Extract the FULL email body
        body = get_email_body(message['payload'])
        
        # Check for attachments
        attachment_names = []
        def find_attachments(payload):
            if 'parts' in payload:
                for part in payload['parts']:
                    if part.get('filename'):
                        attachment_names.append(part['filename'])
                    if 'parts' in part:
                        find_attachments(part)
        
        find_attachments(message['payload'])
        
        # Build comprehensive output
        output_parts = [
            "=== FULL EMAIL CONTENT ===",
            f"**Subject:** {subject}",
            f"**From:** {sender}",
            f"**Date:** {date}"
        ]
        
        if to:
            output_parts.append(f"**To:** {to}")
        if cc:
            output_parts.append(f"**CC:** {cc}")
        if attachment_names:
            output_parts.append(f"**Attachments:** {', '.join(attachment_names)}")
        
        output_parts.append("\n---\n")
        output_parts.append("**Full Body:**")
        output_parts.append(body if body else "(No text content found)")
        output_parts.append("\n=== END EMAIL ===")
        
        return "\n".join(output_parts)
    
    except Exception as e:
        return f"❌ Error fetching email: {str(e)}"


def fetch_email_attachments(query: str) -> str:
    """
    Fetch and parse attachments from recent emails matching the query.
    Use this when user specifically asks about attachments.
    
    Args:
        query: Search query for emails with attachments
    
    Returns:
        Summary of attachments found and saved
    """
    try:
        service = get_gmail_service()
        optimized_query = extract_gmail_query(query) + " has:attachment"
        
        results = service.users().messages().list(
            userId='me',
            q=optimized_query,
            maxResults=2
        ).execute()
        
        messages = results.get('messages', [])
        
        if not messages:
            return f"No emails with attachments found for: '{query}'"
        
        attachment_info = []
        saved_count = 0
        
        for msg_info in messages:
            message = service.users().messages().get(
                userId='me',
                id=msg_info['id'],
                format='full'
            ).execute()
            
            headers = {h['name']: h['value'] for h in message['payload'].get('headers', [])}
            subject = headers.get('Subject', 'No Subject')
            sender = headers.get('From', 'Unknown')
            
            attachments = get_attachments(service, msg_info['id'], message['payload'])
            
            for filename, file_data in attachments[:3]:  # Max 3 per email
                attachment_info.append(f"📎 {filename} (from: {subject})")
        
        if attachment_info:
            return f"Found {len(attachment_info)} attachment(s):\n" + "\n".join(attachment_info) + "\n\n💡 Use `analyze_attachment(email_index, attachment_index)` to read content."
        else:
            return "No attachments found in these emails."
    
    except Exception as e:
        return f"❌ Error: {str(e)}"


# =============================================================================
# TASK 5: THE "BRAIN" (Gemini 2.5 Flash)
# =============================================================================

def get_universal_context() -> dict:
    """
    Get or create the Universal Context.
    Refreshes environment (time) on each call.
    """
    if 'universal_context' not in st.session_state:
        st.session_state.universal_context = initialize_context()
    else:
        # Refresh environment with current time
        ctx = st.session_state.universal_context
        current_time, current_date = get_current_datetime()
        ctx['env']['current_time'] = current_time
        ctx['env']['current_date'] = current_date
        # Refresh todo list from file
        ctx['artifacts']['todo_list'] = load_todo_list()
    
    return st.session_state.universal_context


def build_dynamic_system_prompt() -> str:
    """
    Build system prompt with Universal Context injected.
    """
    context = get_universal_context()
    context_block = format_system_prompt(context)
    
    base_instruction = """You are Emo, a personal AI assistant. Warm, genuine, and smart about when to be brief vs detailed.

CONVERSATION MEMORY:
- You receive [Recent conversation] in your context - USE IT!
- If user asks about something discussed earlier, check the conversation context
- "yes", "yea", "sure", "do it" → refers to your last offer/suggestion in the conversation
- NEVER say "I don't have that information" if it's in [Recent conversation]
- Reference previous messages naturally when relevant

REASONING & LOGIC:
- For puzzles/riddles/logic problems, ALWAYS attempt to solve them
- Never give up with "What would you like me to help with?"
- Show your reasoning step by step
- If a problem has contradictions, explain why it's impossible

CORE PRINCIPLE: Match your response depth to the question depth.
- Simple question → Simple answer
- Complex question → Thoughtful, complete answer
- Asking about content → Give the actual content/summary, not just "yes it exists"

PERSONALITY:
- Natural and conversational - like texting a smart friend
- Genuinely helpful, not performatively helpful
- Warm but not sappy. Friendly but not fake.
- You can be playful, curious, or thoughtful depending on the vibe

FORMATTING STYLE:
- Use plain numbers for lists: 1. 2. 3. NOT emoji numbers like 1️⃣ 2️⃣ 3️⃣
- Minimize emoji usage - only use when it adds real value
- Keep responses clean and professional
- Use markdown formatting: **bold**, *italic*, headings

RESPONSE INTELLIGENCE:

1. GREETINGS & CASUAL:
   - "hey" → "Hey!" or "What's up?" or just dive into relevant context
   - DON'T launch into capabilities or "How can I help?"
   
2. YES/NO + FOLLOW-UP:
   - "Did I get any emails?" → "Yeah, 3 new ones - one from Sarah about the deadline."
   - "Is the meeting tomorrow?" → "Yep, 2pm. Want me to add a reminder?"
   
3. CONTENT QUESTIONS - Give the substance:
   - "What's in that email?" → Actually summarize key points, action items
   - "What about the attachment?" → Describe contents, not just "yes it exists"
   
4. FOLLOW-UP RESPONSES (like "yea", "sure", "do it"):
   - These ALWAYS refer to your last suggestion/offer
   - "Want me to add a reminder?" → "yea" → ADD THE REMINDER with the time you mentioned!
   - Don't ask clarifying questions for info you already stated

5. TASKS & ACTIONS:
   - After doing something → Brief confirmation. "Done." "Added for 2 PM."
   - Include the key detail to confirm: "Added meeting with Sarah at 2pm to your todos."

WHAT TO AVOID:
- Asking for information you already mentioned or found
- "What time?" when you just said "meeting at 2 PM"
- "Who?" when you just said "email from Sarah"
- Over-explaining simple things
- "Sure! I'd be happy to help!" → Just do it
- Emoji numbers (1️⃣ 2️⃣) - use plain 1. 2. 3. instead
- Excessive emojis - keep it minimal
- CALLING THE SAME TOOL REPEATEDLY - if you already searched, use those results!
- Looping on tool calls - ONE search is usually enough

TOOL USAGE RULES:
- Call each tool ONCE per request, then answer with the results
- If search_memory returns nothing, try ONE check_gmail_and_learn, then respond
- If you have results from a tool, USE THEM - don't search again
- For attachments: search_memory("filename") ONCE, if not found, mention it may not be saved yet
- NEVER call the same tool more than twice in one response

NATURAL SPEECH:
- Contractions: "it's", "don't", "I'll", "can't"
- Casual reactions: "Oh nice", "Gotcha", "On it"
- Vary responses - don't start every message the same

CAPABILITIES (use silently):
- Memory: `recall_memory(doc_id)`, `search_memory(query)` - for NON-email content only
- Gmail (ALWAYS fetch fresh - emails are NOT stored in memory!): 
  * `quick_gmail_search(query)` - returns FULL email content, use for ANY email request
  * `get_email_by_index(index)` - use when user says "email 2", "number 3", "the first one"
  * `analyze_attachment(email_idx, att_idx)` - use for "summarize attachment 1", "the second file", "what's in the files"
- Todos: `add_todo`, `get_todos`, `complete_todo`  
- Web: `read_web_page`, `watch_youtube`, `get_news_headlines`
- Session memory: `save_short_term_memory` for temporary context, `query_short_term` to recall
- Quiz: `generate_quiz(quiz_json)` - Create interactive quizzes for learning!

CRITICAL EMAIL RULES (MUST FOLLOW):
- Emails are NOT stored in memory - you MUST fetch fresh from Gmail every time!
- For ANY email request, call `quick_gmail_search` or `get_email_by_index` to get REAL content
- If user says "email 3", "the second one", "number 1" -> YOU MUST USE `get_email_by_index(index)`
- DO NOT use `quick_gmail_search` for "the third one" - it will fail! Use `get_email_by_index(3)`
- NEVER generate/guess email content - always fetch fresh from Gmail
- The tool returns the COMPLETE email body - display it exactly as returned

⚠️ EMAIL ANTI-HALLUCINATION (CRITICAL - READ THIS):
- When you receive [get_email_by_index result] or [quick_gmail_search result], that IS the real email
- COPY the Subject, From, Date, Body EXACTLY from the tool result - do not modify
- NEVER invent sender names, dates, meeting times, or body content
- If tool says "From: Sarah Johnson", display "From: Sarah Johnson" - not some other name
- If tool says meeting is at "2:00 PM", say "2:00 PM" - not "10:00 AM"
- The tool result is THE TRUTH - your job is to present it, not rewrite it
- DO NOT add "plausible" content that is not in the tool result - this is HALLUCINATION

ATTACHMENT RULES (CRITICAL - DO NOT VIOLATE):
- You CANNOT "access" or "read" attachment files directly - you MUST call `analyze_attachment`
- When user asks "summarize the attachment", "what's in the file", "show me the PDF" -> CALL THE TOOL
- NEVER say "I can't access the file" - you CAN access it via `analyze_attachment`!
- NEVER generate "plausible" or "typical" content for files - this is HALLUCINATION
- If you don't have tool output for a file -> call `analyze_attachment(email_idx, att_idx)`
- ONLY present file content that comes from the `analyze_attachment` tool result

QUIZ GENERATION - When user asks for a quiz:
- Use `generate_quiz` tool with a JSON string containing the quiz
- Create 3-5 questions based on the content they want to be tested on
- ONLY use these 2 question types (auto-gradable):
  * multiple_choice - always include 4 options
  * true_false - simple true/false statements
- Always include explanations for each answer
- For MATH quizzes: USE LATEX! Write equations with $...$ syntax:
  * Question: "What is the derivative of $x^2$?"
  * Options: ["$2x$", "$x$", "$2x^2$", "$x^2$"]
  * Explanation: "Using power rule: $\\frac{d}{dx}x^n = nx^{n-1}$, so $\\frac{d}{dx}x^2 = 2x$"
- JSON format example:
  {"title": "Topic Quiz", "questions": [
    {"id": 1, "type": "multiple_choice", "question": "What is $\\sqrt{16}$?", "options": ["$2$","$4$","$8$","$16$"], "correct": 1, "explanation": "Since $4 \\times 4 = 16$, $\\sqrt{16} = 4$"},
    {"id": 2, "type": "true_false", "question": "$x^2 + y^2 = r^2$ is the equation of a circle", "correct": true, "explanation": "This is the standard form of a circle centered at origin"},
    {"id": 3, "type": "multiple_choice", "question": "Solve: $2x + 4 = 10$", "options": ["$x=2$","$x=3$","$x=4$","$x=5$"], "correct": 1, "explanation": "$2x = 6$, so $x = 3$"}
  ]}
- The quiz renders LaTeX beautifully for math content!

CRITICAL - DO NOT HALLUCINATE CAPABILITIES:
- You can ONLY do what's listed above. NOTHING ELSE.
- You CANNOT: open links, open browsers, send emails, block emails, unsubscribe, click anything
- You CANNOT: access user's computer, run programs, install apps, make calls
- You CANNOT: schedule meetings in calendar apps, set system alarms, control devices
- NEVER offer options like "I can open this in your browser" - YOU CANNOT DO THAT
- NEVER say "I can block these emails" - YOU CANNOT DO THAT
- If user asks for something you can't do, say "I can't do that directly, but here's what I found..."
- Only offer actions from your ACTUAL tool list above

CRITICAL - USE TOOL RESULTS:
- When you see [watch_youtube result]: - YOU HAVE THE FULL TRANSCRIPT with timestamps!
- The transcript has timestamps like [00:00], [01:00], [22:00] etc.
- If user asks "what happens at minute 22", LOOK at the [22:00] section in the transcript!
- NEVER say "I can't access the transcript" when it's RIGHT THERE in your context
- NEVER say "I don't have direct access to YouTube's internal video data" - you DO have the transcript!
- When you see [read_web_page result]: - YOU HAVE THE PAGE CONTENT, use it!
- When you see [quick_gmail_search result]: - YOU HAVE THE EMAILS, reference them!
- Tool results in your context ARE your data source - USE THEM DIRECTLY
- If the transcript shows [22:00] followed by text, that IS minute 22 - quote it!

PROACTIVE:
- Find deadlines → add to todos WITH the time
- User shares facts → save to long_term_memory quietly

DEADLINE REMINDERS (when you see [DEADLINE ALERTS]):
- URGENT (< 15 min): Interrupt. "Hey, your call is in 10 min."
- SOON (15-30 min): Mention naturally. "That deadline's coming up."
- UPCOMING (30-120 min): Casual. "You've got that thing at 11."

ATTACHMENTS - YOU CAN READ THEM:
- PDF, DOCX, XLSX content is in memory after Gmail fetch
- Search: `search_memory("filename.pdf")`
- NEVER say "I can't read attachments"

Remember: You have full conversation context. Use it. Don't ask for what you know.

MATH & LATEX FORMATTING:
The UI renders LaTeX equations beautifully! Use it for any math.

SYNTAX RULES:
- Use $...$ for inline math within text: "where $a \\neq 0$"
- Use $$...$$ on its OWN LINE for display equations

EXAMPLE - inline math in sentence:
The quadratic formula gives us $x = \\frac{-b \\pm \\sqrt{b^2-4ac}}{2a}$ where $a$, $b$, $c$ are coefficients.

EXAMPLE - display equation (must be on own line):

$$x = \\frac{-b \\pm \\sqrt{b^2-4ac}}{2a}$$

COMMON LATEX:
- Fractions: $\\frac{a}{b}$
- Roots: $\\sqrt{x}$, $\\sqrt[n]{x}$
- Powers: $x^2$, $x^{n+1}$
- Subscripts: $x_n$, $a_{i,j}$
- Greek: $\\alpha$, $\\beta$, $\\pi$, $\\Delta$
- Operators: $\\sum$, $\\int$, $\\lim$, $\\prod$
- Relations: $\\geq$, $\\leq$, $\\neq$, $\\approx$

NEVER USE \\[...\\] or \\(...\\) syntax! Only use $ and $$.

TEXT FORMATTING:
- **bold** for emphasis
- *italic* for terms
- `code` for code
- ## Headings for structure
- Use proper structure for complex explanations"""
    
    return f"{base_instruction}\n\n{context_block}"


def get_openai_tools():
    """
    Build OpenAI-compatible tools schema for Groq API.
    """
    return [
        {
            "type": "function",
            "function": {
                "name": "check_gmail_and_learn",
                "description": "Search Gmail and save emails to memory. Fast by default (skips attachments). Use for finding email content.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Natural language search query (e.g., 'find my upcoming appointments')"}
                    },
                    "required": ["query"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "quick_gmail_search",
                "description": "Search Gmail and return FULL email content including complete body, attachments list, all details. Use this for any email request - 'check emails', 'show email', 'give me email content', etc.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Natural language search query"},
                        "max_results": {"type": "integer", "description": "Maximum number of results (default: 3)", "default": 3}
                    },
                    "required": ["query"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "get_email_by_index",
                "description": "Get full content of an email by its index number (e.g., 'email 2', 'number 3') from the LAST search results. Use this when user refers to an email by number.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "index": {"type": "integer", "description": "The index number of the email (1, 2, 3...)"}
                    },
                    "required": ["index"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "analyze_attachment",
                "description": "Analyze an attachment from a cached email and return a 'cool' summary (File DNA). Use when user asks to 'summarize attachment', 'what is in the file', etc.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "email_index": {"type": "integer", "description": "The index of the email (1, 2, 3...)"},
                        "attachment_index": {"type": "integer", "description": "The index of the attachment (default: 1)", "default": 1}
                    },
                    "required": ["email_index"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "fetch_email_attachments",
                "description": "Fetch and parse attachments from emails. Use ONLY when user specifically asks about attachments or documents in emails.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Search query to find emails with attachments"}
                    },
                    "required": ["query"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "get_full_email",
                "description": "Fetch the FULL content of an email. Use this when the user wants complete email details, full body text, all attendees, meeting links, agenda items, etc. - NOT just a summary or preview.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Search query to find the specific email (e.g., 'Weekly Standup from Trieu An')"}
                    },
                    "required": ["query"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "add_todo",
                "description": "Add a new task to the todo list. Include time/deadline if mentioned.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "task": {"type": "string", "description": "The task description including any deadline/time"}
                    },
                    "required": ["task"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "get_todos",
                "description": "Get all tasks from the todo list.",
                "parameters": {"type": "object", "properties": {}}
            }
        },
        {
            "type": "function",
            "function": {
                "name": "complete_todo",
                "description": "Mark a task as complete by its ID.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "task_id": {"type": "integer", "description": "The ID of the task to complete"}
                    },
                    "required": ["task_id"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "read_web_page",
                "description": "Fetch and return content from a web page URL.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "url": {"type": "string", "description": "The URL of the web page to read"}
                    },
                    "required": ["url"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "watch_youtube",
                "description": "Get transcript from a YouTube video.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "video_url": {"type": "string", "description": "The YouTube video URL"}
                    },
                    "required": ["video_url"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "get_news_headlines",
                "description": "Get current news headlines, optionally filtered by topic.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "topic": {"type": "string", "description": "Optional topic to filter news"}
                    }
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "recall_memory",
                "description": "Retrieve full content of a specific memory by its document ID.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "doc_id": {"type": "string", "description": "The document ID to recall"}
                    },
                    "required": ["doc_id"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "search_memory",
                "description": "Search through stored memories using semantic search.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "What to search for in memory"}
                    },
                    "required": ["query"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "save_short_term_memory",
                "description": "Save temporary context for this session only.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "content": {"type": "string", "description": "The information to remember"},
                        "context": {"type": "string", "description": "Why this is relevant"},
                        "importance": {"type": "string", "enum": ["low", "normal", "high"], "default": "normal"}
                    },
                    "required": ["content", "context"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "save_long_term_memory",
                "description": "Save permanent personal information (preferences, facts about user).",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "fact": {"type": "string", "description": "The personal fact/preference to remember"},
                        "category": {"type": "string", "enum": ["identity", "preference", "relationship", "date", "skill", "other"]}
                    },
                    "required": ["fact", "category"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "query_short_term",
                "description": "Search short-term memory for context from this session.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "What to search for"}
                    },
                    "required": ["query"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "query_long_term",
                "description": "Search long-term memory for permanent personal facts.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "What personal information to search for"}
                    },
                    "required": ["query"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "generate_quiz",
                "description": "Generate an interactive quiz for learning. Create questions based on content the user wants to be tested on. The quiz will be displayed interactively in the UI.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "quiz_json": {
                            "type": "string",
                            "description": """JSON string with quiz structure. Format:
{
  "title": "Quiz Title",
  "description": "Optional description",
  "questions": [
    {"id": 1, "type": "multiple_choice", "question": "Question text?", "options": ["A", "B", "C", "D"], "correct": 0, "explanation": "Why this is correct"},
    {"id": 2, "type": "true_false", "question": "Statement to evaluate", "correct": true, "explanation": "Explanation"},
    {"id": 3, "type": "short_answer", "question": "Question?", "correct": "Answer", "explanation": "Explanation"}
  ]
}
Types: multiple_choice (options + correct index 0-3), true_false (correct is true/false), short_answer (correct is text)"""
                        }
                    },
                    "required": ["quiz_json"]
                }
            }
        }
    ]


def setup_tools():
    """Setup available tools mapping."""
    st.session_state.available_tools = {
        'check_gmail_and_learn': check_gmail_and_learn,
        'quick_gmail_search': quick_gmail_search,
        'get_email_by_index': get_email_by_index,
        'analyze_attachment': analyze_attachment,
        'fetch_email_attachments': fetch_email_attachments,
        'get_full_email': get_full_email,
        'add_todo': add_todo,
        'get_todos': get_todos,
        'complete_todo': complete_todo,
        'read_web_page': read_web_page,
        'watch_youtube': watch_youtube,
        'get_news_headlines': get_news_headlines,
        'recall_memory': recall_memory,
        'search_memory': search_memory,
        'save_short_term_memory': save_short_term_memory,
        'save_long_term_memory': save_long_term_memory,
        'save_project_memory': save_project_memory,
        'query_short_term': query_short_term,
        'query_long_term': query_long_term,
        'query_project': query_project,
        'list_all_projects': list_all_projects,
        'update_long_term_memory': update_long_term_memory,
        'generate_quiz': generate_quiz,  # Interactive quiz generator
    }


def get_chat_messages():
    """
    Get or create chat message history for Groq.
    """
    if 'chat_messages' not in st.session_state:
        st.session_state.chat_messages = []
    return st.session_state.chat_messages


# =============================================================================
# TOOL EXECUTION DISPLAY SYSTEM
# =============================================================================

def format_tool_args(args: dict) -> str:
    """Format tool arguments for display."""
    if not args:
        return "none"
    
    formatted = []
    for key, value in args.items():
        if isinstance(value, str) and len(value) > 50:
            value = value[:47] + "..."
        formatted.append(f"{key}={repr(value)}")
    return ", ".join(formatted)


def execute_tool_with_display(tool_name: str, tool_args: dict, tool_container) -> str:
    """
    Execute a tool and display real-time status.
    
    Args:
        tool_name: Name of the tool to execute
        tool_args: Arguments to pass to the tool
        tool_container: Streamlit container for displaying status
    
    Returns:
        Tool execution result as string
    """
    available_tools = st.session_state.get('available_tools', {})
    
    if tool_name not in available_tools:
        return f"Error: Unknown tool '{tool_name}'"
    
    tool_func = available_tools[tool_name]
    
    # Execute the tool
    start_time = time.time()
    try:
        result = tool_func(**tool_args)
        execution_time = time.time() - start_time
        
        # Update status to complete (minimal display)
        with tool_container:
            st.markdown(f"""
<div style="font-family: monospace; font-size: 0.75rem; padding: 0.5rem; background: #0d1117; border-radius: 4px; margin: 0.25rem 0; border-left: 2px solid #3fb950;">
<span style="color: #3fb950;">done</span> · <span style="color: #8b949e;">{tool_name}</span> · <span style="color: #666;">{execution_time:.2f}s</span>
</div>
""", unsafe_allow_html=True)
        
        return str(result)
    
    except Exception as e:
        execution_time = time.time() - start_time
        error_msg = str(e)[:100]
        
        with tool_container:
            st.markdown(f"""
<div style="font-family: monospace; font-size: 0.75rem; padding: 0.5rem; background: #0d1117; border-radius: 4px; margin: 0.25rem 0; border-left: 2px solid #f85149;">
<span style="color: #f85149;">error</span> · <span style="color: #8b949e;">{tool_name}</span>
</div>
""", unsafe_allow_html=True)
        
        return f"Error: {error_msg}"


def display_tool_call_start(tool_name: str, tool_args: dict, tool_container, index: int):
    """Display a tool call in progress."""
    with tool_container:
        st.markdown(f"""
<div style="font-family: monospace; font-size: 0.75rem; padding: 0.5rem; background: #0d1117; border-radius: 4px; margin: 0.25rem 0; border-left: 2px solid #58a6ff;">
<span style="color: #58a6ff;">running</span> · <span style="color: #8b949e;">{tool_name}</span>
</div>
""", unsafe_allow_html=True)


def display_thinking(thinking_container, query_type: str, tools_needed: list, model: str, tool_results: list = None):
    """
    Display AI thinking process in a collapsible section - like Gemini/Claude.
    Shows the reasoning steps the AI is taking.
    """
    with thinking_container:
        # Build thinking steps
        thinking_steps = []
        
        # Step 1: Understanding
        thinking_steps.append(f"Understanding the request: analyzing query type...")
        
        # Step 2: Classification result
        if query_type == 'simple_chat':
            thinking_steps.append("This is a casual/greeting message - no tools needed.")
        elif query_type == 'email':
            thinking_steps.append("User is asking about emails - need to search Gmail.")
        elif query_type == 'todo':
            thinking_steps.append("This is about tasks/todos - checking task manager.")
        elif query_type == 'web':
            thinking_steps.append("User wants web content - fetching from internet.")
        elif query_type == 'memory':
            thinking_steps.append("Looking for stored information in memory.")
        else:
            thinking_steps.append(f"Query classified as: {query_type}")
        
        # Step 3: Tools
        if tools_needed:
            tools_str = ", ".join(tools_needed)
            thinking_steps.append(f"Tools to use: {tools_str}")
        
        # Step 4: Model selection
        model_name = model.split('/')[-1] if '/' in model else model
        thinking_steps.append(f"Using model: {model_name}")
        
        # Step 5: Tool results summary
        if tool_results:
            for tr in tool_results:
                result_preview = tr.get('result', '')[:100]
                if result_preview:
                    thinking_steps.append(f"Got results from {tr.get('tool', 'tool')}")
        
        # Render as collapsible expander with nice styling
        with st.expander("💭 Thinking...", expanded=False):
            for i, step in enumerate(thinking_steps, 1):
                st.markdown(f"""
<div style="font-size: 0.85rem; color: #8b949e; padding: 0.25rem 0; border-left: 2px solid #30363d; padding-left: 0.75rem; margin: 0.25rem 0;">
{step}
</div>
""", unsafe_allow_html=True)


def chat_with_emo(user_message: str, memory_context: str = "", show_thinking: bool = False, tool_display_container=None) -> dict:
    """
    Smart 3-phase chat architecture for smaller models:
    1. CLASSIFY: Fast keyword-based routing (no AI)
    2. EXECUTE: Run needed tools directly
    3. RESPOND: Single AI call with all context
    
    This is much faster and more reliable than letting the AI decide tools.
    """
    # Setup tools if not already done
    if 'available_tools' not in st.session_state:
        setup_tools()
    
    client = get_groq_client()
    messages = get_chat_messages()
    
    result = {
        'response': '',
        'thinking': '',
        'tools_used': [],
        'tool_results': [],
        'tool_executions': []
    }
    
    # Create thinking container if show_thinking is enabled
    thinking_container = None
    if show_thinking and tool_display_container:
        thinking_container = tool_display_container
    
    try:
        # ========== PHASE 1: CLASSIFY (no AI needed) ==========
        query_type, tools_needed = classify_query(user_message)
        print(f"[Emo] Query type: {query_type}, Tools: {tools_needed}")
        
        # Select best model for this task
        model = select_model_for_task(query_type, len(tools_needed) > 0)
        print(f"[Emo] Using model: {model}")
        
        # ========== PHASE 2: EXECUTE TOOLS ==========
        tool_context = []
        
        for tool_name in tools_needed:
            if tool_name not in st.session_state.available_tools:
                continue
            
            # Prepare tool arguments based on the query
            tool_args = prepare_tool_args(tool_name, user_message)
            
            result['tools_used'].append(tool_name)
            
            # Display tool start
            if tool_display_container:
                display_tool_call_start(tool_name, tool_args, tool_display_container, 0)
            
            # Execute tool
            start_time = time.time()
            tool_result = execute_tool_with_display(
                tool_name,
                tool_args,
                tool_display_container if tool_display_container else st.empty()
            )
            duration = time.time() - start_time
            
            result['tool_executions'].append({
                'tool': tool_name,
                'args': tool_args,
                'duration': duration,
                'result': tool_result[:300] if tool_result else ""
            })
            
            # Add to context for AI
            tool_context.append(f"[{tool_name} result]:\n{tool_result}")
        
        # Display thinking process (like Gemini/Claude)
        if show_thinking and thinking_container:
            display_thinking(thinking_container, query_type, tools_needed, model, result['tool_executions'])
        
        # ========== PHASE 3: GENERATE RESPONSE ==========
        # Build context
        context_parts = []
        
        # Add recent conversation summary for multi-turn context
        # This ensures the AI remembers what was discussed even without memory search
        if messages and len(messages) > 0:
            recent_msgs = messages[-8:]  # Last 8 messages
            conv_summary = []
            for msg in recent_msgs:
                role = "User" if msg['role'] == 'user' else "Emo"
                content = msg['content'][:300] + "..." if len(msg['content']) > 300 else msg['content']
                conv_summary.append(f"{role}: {content}")
            if conv_summary:
                context_parts.append(f"[Recent conversation]:\n" + "\n".join(conv_summary))
        
        # Add memory context if available
        if memory_context:
            context_parts.append(f"[Relevant memories]:\n{memory_context}")
        
        # Add tool results
        if tool_context:
            context_parts.append("\n".join(tool_context))
        
        # Add deadline reminders
        reminders = get_smart_reminders()
        if reminders:
            context_parts.append(reminders)
        
        # Build the full prompt
        # Detect if this is a reasoning/logic problem
        is_logic_problem = any(kw in user_message.lower() for kw in [
            'puzzle', 'riddle', 'logic', 'brothers', 'sisters', 'sibling',
            'if then', 'how many', 'figure out', 'solve this', 'brain teaser',
            'contradiction', 'paradox'
        ])
        
        if context_parts:
            # Check if we have YouTube transcript in the context
            has_youtube = any('[watch_youtube result]' in part for part in context_parts)
            
            if has_youtube:
                full_prompt = f"""Here is the context and tool results:

{chr(10).join(context_parts)}

User message: {user_message}

IMPORTANT: You have the FULL YouTube transcript above with timestamps like [00:00], [01:00], [22:00], etc.
If user asks about a specific minute, LOOK at that timestamp section and quote the content directly.
Do NOT say you can't access the transcript - it's RIGHT ABOVE in the [watch_youtube result] section!"""
            else:
                # Check if we have email tool results that need anti-hallucination enforcement
                has_email_result = any(
                    '[get_email_by_index result]' in part or 
                    '[quick_gmail_search result]' in part or
                    'EMAIL #' in part or
                    '=== FULL CONTENT FOR EMAIL' in part
                    for part in context_parts
                )
                
                if has_email_result:
                    full_prompt = f"""Here is the context and tool results:

{chr(10).join(context_parts)}

User message: {user_message}

⚠️ CRITICAL ANTI-HALLUCINATION RULES FOR EMAIL CONTENT ⚠️

1. The email content from the tool result above is the REAL email from Gmail.
2. You MUST display this email content EXACTLY as shown - do NOT modify, summarize, or rewrite it.
3. DO NOT generate your own version of the email - use ONLY what is in the tool result.
4. Copy the exact Subject, From, Date, Body, and Attachments from the tool result.
5. If the tool result shows specific content, THAT is the truth - nothing else.

Present the email content from the tool result to the user. Do not make up different content."""
                else:
                    full_prompt = f"""Here is the context and tool results:

{chr(10).join(context_parts)}

User message: {user_message}

Based on the above context and results, provide a helpful response."""
        else:
            full_prompt = user_message
        
        # Add reasoning instructions for logic problems
        if is_logic_problem:
            full_prompt += """\n\nThis appears to be a logic/reasoning problem. Please:
1. Identify all the given facts and constraints
2. Look for any contradictions or inconsistencies
3. If contradictory, explain WHY it's impossible
4. If solvable, show step-by-step reasoning
5. State your final answer clearly

Do NOT give up - always attempt to analyze the problem even if it seems tricky."""
        
        # Note: Don't add thinking instructions - Qwen already outputs <think> blocks
        # We'll parse them out automatically if present
        
        # Add to message history
        messages.append({"role": "user", "content": user_message})
        
        # Build system prompt (shorter for simple queries)
        # Keep thinking enabled (helps quality) but we'll parse it out for display
        if query_type == 'simple_chat':
            system_prompt = """You are Emo, a friendly AI assistant. Be warm, brief, and natural.

You have access to the recent conversation history. Use it to understand context.
If user says "yea", "yes", "sure", "do it" - they are confirming your last offer/suggestion."""
        else:
            system_prompt = build_dynamic_system_prompt()
        
        # Single AI call to generate response
        # Include recent conversation history for context
        api_messages = [
            {"role": "system", "content": system_prompt},
        ]
        
        # Add recent conversation history (last 6 messages for context)
        recent_history = messages[-6:] if len(messages) > 6 else messages
        for msg in recent_history:
            api_messages.append({"role": msg["role"], "content": msg["content"]})
        
        # Add current user message with context
        api_messages.append({"role": "user", "content": full_prompt})
        
        # For quiz mode, ask AI to generate JSON directly (no tool calling - more reliable)
        if query_type == 'quiz':
            quiz_prompt = full_prompt + """

Generate an interactive quiz as a JSON code block.
ONLY use these question types (instant auto-grading, no typing needed):

1. **multiple_choice** - User clicks one option
2. **true_false** - User clicks True or False

```json
{
  "title": "Quiz Title",
  "description": "Brief description",
  "questions": [
    {
      "id": 1,
      "type": "multiple_choice",
      "question": "What is the capital of France?",
      "options": ["London", "Paris", "Berlin", "Madrid"],
      "correct": 1,
      "explanation": "Paris is the capital of France"
    },
    {
      "id": 2,
      "type": "true_false",
      "question": "Python is a compiled language",
      "correct": false,
      "explanation": "Python is an interpreted language"
    },
    {
      "id": 3,
      "type": "multiple_choice",
      "question": "Which is NOT a programming language?",
      "options": ["Java", "HTML", "Python", "C++"],
      "correct": 1,
      "explanation": "HTML is a markup language, not a programming language"
    }
  ]
}
```

Rules:
- "correct" for multiple_choice = index of correct option (0, 1, 2, or 3)
- "correct" for true_false = true or false (boolean)
- Always include "explanation" for learning
- Create 5-6 questions mixing both types
- Make questions relevant to the content

Output ONLY the JSON block, no other text."""

            api_messages[-1] = {"role": "user", "content": quiz_prompt}
            
            response = client.chat.completions.create(
                model=model,
                messages=api_messages,
                temperature=0.7,
                max_tokens=2048,
            )
            
            raw_response = response.choices[0].message.content or ""
            print(f"[Quiz] Raw response length: {len(raw_response)}")
            
            # Extract JSON from response - try multiple patterns
            import re
            quiz_json_str = None
            
            # Pattern 1: ```json ... ``` (greedy match)
            json_match = re.search(r'```json\s*(.*?)\s*```', raw_response, re.DOTALL)
            if json_match:
                quiz_json_str = json_match.group(1)
                print(f"[Quiz] Found JSON in code block, length: {len(quiz_json_str)}")
            
            # Pattern 2: ``` ... ``` without json tag
            if not quiz_json_str:
                json_match = re.search(r'```\s*(\{.*?"questions".*?\})\s*```', raw_response, re.DOTALL)
                if json_match:
                    quiz_json_str = json_match.group(1)
                    print(f"[Quiz] Found JSON in generic code block")
            
            # Pattern 3: Raw JSON object with questions
            if not quiz_json_str:
                json_match = re.search(r'(\{[^{}]*"questions"\s*:\s*\[.*?\]\s*\})', raw_response, re.DOTALL)
                if json_match:
                    quiz_json_str = json_match.group(1)
                    print(f"[Quiz] Found raw JSON object")
            
            # Pattern 4: Find opening { and match to closing }
            if not quiz_json_str:
                start = raw_response.find('{"title"')
                if start == -1:
                    start = raw_response.find('{\n')
                if start >= 0:
                    # Count braces to find matching end
                    depth = 0
                    end = start
                    for i, c in enumerate(raw_response[start:]):
                        if c == '{':
                            depth += 1
                        elif c == '}':
                            depth -= 1
                            if depth == 0:
                                end = start + i + 1
                                break
                    quiz_json_str = raw_response[start:end]
                    print(f"[Quiz] Extracted JSON by brace matching")
            
            if quiz_json_str:
                # Clean up the JSON string
                quiz_json_str = quiz_json_str.strip()
                
                # Display tool execution
                if tool_display_container:
                    display_tool_call_start('generate_quiz', {'questions': 'parsing...'}, tool_display_container, 0)
                
                quiz_result = generate_quiz(quiz_json_str)
                result['tools_used'].append('generate_quiz')
                result['response'] = quiz_result
                
                # Add to history
                messages.append({"role": "assistant", "content": result['response']})
                st.session_state.chat_messages = messages
                return result
            else:
                # Fallback - show raw response
                raw_response = "I couldn't generate a proper quiz format. Let me try a simpler approach:\n\n" + raw_response
        else:
            response = client.chat.completions.create(
                model=model,
                messages=api_messages,
                temperature=0.7,
                max_tokens=1024,
            )
            raw_response = response.choices[0].message.content or ""
        
        # Parse thinking from response if present (handle multiple formats)
        import re
        # Check for <think> blocks (Qwen outputs these)
        # Handle both closed and unclosed think blocks
        if '<think>' in raw_response.lower():
            # Try closed block first
            think_pattern = r'<think>\s*(.*?)\s*</think>'
            think_matches = re.findall(think_pattern, raw_response, re.DOTALL | re.IGNORECASE)
            
            if think_matches:
                result['thinking'] = '\n'.join(think_matches).strip()
                cleaned = re.sub(think_pattern, '', raw_response, flags=re.DOTALL | re.IGNORECASE).strip()
                result['response'] = cleaned
            else:
                # Handle unclosed <think> block (everything after <think> is thinking)
                parts = re.split(r'<think>', raw_response, flags=re.IGNORECASE)
                if len(parts) > 1:
                    # First part is the response (if any), rest is thinking
                    result['response'] = parts[0].strip()
                    result['thinking'] = '\n'.join(parts[1:]).strip()
                    # If response is empty, it means everything was thinking - just hide it
                    if not result['response']:
                        result['response'] = "I understand. What would you like me to help with?"
                else:
                    result['response'] = raw_response
        else:
            result['response'] = raw_response
        
        # Final cleanup - remove any stray think tags
        result['response'] = re.sub(r'</?think>', '', result['response'], flags=re.IGNORECASE).strip()
        
        # Add to history (without thinking block)
        messages.append({"role": "assistant", "content": result['response']})
        st.session_state.chat_messages = messages
        
        return result
    
    except Exception as e:
        print(f"[Emo Error] {str(e)}")
        result['response'] = f"Sorry, I encountered an error: {str(e)}"
        return result


def prepare_tool_args(tool_name: str, user_message: str) -> dict:
    """
    Prepare arguments for a tool based on the user message.
    No AI needed - just extract relevant info.
    """
    msg_lower = user_message.lower()
    
    if tool_name in ['quick_gmail_search', 'check_gmail_and_learn', 'fetch_email_attachments', 'get_full_email']:
        # Extract email search query
        query = user_message
        # Clean up common prefixes
        for prefix in ['check', 'find', 'get', 'show', 'give me', 'search for']:
            if msg_lower.startswith(prefix):
                query = user_message[len(prefix):].strip()
                break
        return {'query': query}
    
    elif tool_name == 'get_email_by_index':
        # Extract index number (digits or words)
        import re
        
        # Check for digits first
        numbers = re.findall(r'\d+', user_message)
        if numbers:
            return {'index': int(numbers[0])}
            
        # Check for ordinal words
        msg_lower = user_message.lower()
        word_map = {
            'first': 1, '1st': 1, 'one': 1,
            'second': 2, '2nd': 2, 'two': 2,
            'third': 3, '3rd': 3, 'three': 3,
            'fourth': 4, '4th': 4, 'four': 4,
            'fifth': 5, '5th': 5, 'five': 5
        }
        
        for word, index in word_map.items():
            if word in msg_lower:
                return {'index': index}
                
        return {'index': 1}  # Default to first email
        
    elif tool_name == 'analyze_attachment':
        # Extract email index and attachment index
        import re
        numbers = re.findall(r'\d+', user_message)
        
        # Default to last viewed email if available, else 1
        email_idx = st.session_state.get('last_viewed_email_index', 1)
        att_idx = 1
        
        # Check for plural "attachments" or "all" -> analyze ALL (index 0)
        if "attachments" in user_message.lower() or "all files" in user_message.lower() or "all attachment" in user_message.lower() or ("all" in user_message.lower() and "attachment" in user_message.lower()):
            att_idx = 0
            
        # Check for ordinal words for attachment index
        msg_lower = user_message.lower()
        word_map = {
            'first': 1, '1st': 1, 'one': 1,
            'second': 2, '2nd': 2, 'two': 2,
            'third': 3, '3rd': 3, 'three': 3,
            'fourth': 4, '4th': 4, 'four': 4,
            'fifth': 5, '5th': 5, 'five': 5
        }
        
        # If explicit digits found
        if len(numbers) >= 2:
            # "attachment 2 from email 3"
            att_idx = int(numbers[0])
            email_idx = int(numbers[1])
        elif len(numbers) == 1:
            # "summarize attachment 2"
            att_idx = int(numbers[0])
        else:
            # No digits, check for words
            for word, index in word_map.items():
                if word in msg_lower:
                    att_idx = index
                    break
            
        return {'email_index': email_idx, 'attachment_index': att_idx}
    
    elif tool_name == 'add_todo':
        # Extract task from message
        task = user_message
        for prefix in ['add', 'create', 'remind me to', 'add todo', 'add task']:
            if msg_lower.startswith(prefix):
                task = user_message[len(prefix):].strip()
                break
        return {'task': task}
    
    elif tool_name == 'complete_todo':
        # Extract task number or name
        import re
        numbers = re.findall(r'\d+', user_message)
        if numbers:
            return {'task_id': int(numbers[0])}
        return {'task_id': 1}  # Default to first task
    
    elif tool_name == 'get_todos':
        return {}
    
    elif tool_name == 'read_web_page':
        # Extract URL
        import re
        urls = re.findall(r'https?://\S+', user_message)
        if urls:
            return {'url': urls[0]}
        return {'url': user_message}
    
    elif tool_name == 'watch_youtube':
        # Extract YouTube URL - parameter name is 'video_url'
        import re
        # Match full URLs
        full_urls = re.findall(r'https?://(?:www\.)?(?:youtube\.com/watch\?v=|youtu\.be/)[\w-]+[^\s]*', user_message)
        if full_urls:
            return {'video_url': full_urls[0]}
        # Match partial URLs
        partial_urls = re.findall(r'(?:youtube\.com/watch\?v=|youtu\.be/)[\w-]+', user_message)
        if partial_urls:
            url = partial_urls[0]
            if not url.startswith('http'):
                url = 'https://' + url
            return {'video_url': url}
        # If user just pasted a URL, use it directly
        if 'youtu' in user_message.lower():
            return {'video_url': user_message.strip()}
        return {'video_url': user_message}
    
    elif tool_name == 'get_news_headlines':
        # Map topics to appropriate news URLs
        topic_urls = {
            'ai': 'https://techcrunch.com/category/artificial-intelligence/',
            'artificial intelligence': 'https://techcrunch.com/category/artificial-intelligence/',
            'tech': 'https://techcrunch.com/',
            'technology': 'https://techcrunch.com/',
            'world': 'https://www.bbc.com/news/world',
            'business': 'https://www.bbc.com/news/business',
            'science': 'https://www.bbc.com/news/science_and_environment',
            'sports': 'https://www.bbc.com/sport',
            'vietnam': 'https://vnexpress.net/',
            'vn': 'https://vnexpress.net/',
        }
        
        # Extract topic and find matching URL
        topic = None
        for kw in ['about', 'on', 'regarding', 'for']:
            if kw in msg_lower:
                idx = msg_lower.index(kw)
                topic = user_message[idx + len(kw):].strip().lower()
                # Remove trailing punctuation
                topic = topic.rstrip('?!.')
                break
        
        # Find URL for topic
        url = 'https://news.google.com/'  # Default
        if topic:
            for key, topic_url in topic_urls.items():
                if key in topic:
                    url = topic_url
                    break
        
        return {'url': url}
    
    elif tool_name == 'search_memory':
        return {'query': user_message}
    
    elif tool_name == 'save_long_term_memory':
        # Extract what to save
        content = user_message
        for prefix in ['save', 'remember', 'note']:
            if msg_lower.startswith(prefix):
                content = user_message[len(prefix):].strip()
                break
        return {'content': content, 'category': 'user_note'}
    
    # Default: pass message as query
    return {'query': user_message}


# =============================================================================
# TASK 6: THE FRONTEND (Streamlit)
# =============================================================================

def get_custom_css():
    """Return custom CSS for clean, minimal UI."""
    return """
    <style>
        /* Hide streamlit defaults */
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        
        /* Keep sidebar toggle visible */
        [data-testid="collapsedControl"] {
            display: flex !important;
            visibility: visible !important;
            color: #b0b0b0 !important;
        }
        
        /* Compact sidebar */
        [data-testid="stSidebar"] {
            background-color: #1a1a1a;
            border-right: 1px solid #2a2a2a;
        }
        
        [data-testid="stSidebar"] > div:first-child {
            padding-top: 0.5rem;
            padding-bottom: 60px;  /* Space for fixed bottom buttons */
        }
        
        /* Compact chat buttons */
        [data-testid="stSidebar"] button {
            padding: 0.4rem 0.6rem !important;
            font-size: 0.85rem !important;
            margin: 1px 0 !important;
        }
        
        /* Section labels - smaller */
        .section-label {
            font-size: 0.7rem;
            font-weight: 500;
            color: #555;
            padding: 0.25rem 0;
            margin: 0;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }
        
        /* Remove extra padding from columns */
        [data-testid="stSidebar"] [data-testid="column"] {
            padding: 0 !important;
        }
        
        /* Compact text input */
        [data-testid="stSidebar"] input {
            font-size: 0.85rem !important;
            padding: 0.4rem !important;
        }
        
        /* Compact dividers */
        [data-testid="stSidebar"] hr {
            margin: 0.5rem 0 !important;
        }
    </style>
    """


def initialize_session_state():
    """
    Initialize Streamlit session state variables.
    """
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    
    if 'gmail_authenticated' not in st.session_state:
        st.session_state.gmail_authenticated = False
    
    # Initialize Universal Context
    if 'universal_context' not in st.session_state:
        st.session_state.universal_context = initialize_context()
    
    # Initialize Chat History session
    if 'current_session_id' not in st.session_state:
        st.session_state.current_session_id = None
    
    # Set current session ID for short-term memory scoping
    if st.session_state.current_session_id:
        set_current_session_id(st.session_state.current_session_id)
    
    # AI Thinking mode - default ON
    if 'show_thinking' not in st.session_state:
        st.session_state.show_thinking = True
    
    # UI state
    if 'show_settings' not in st.session_state:
        st.session_state.show_settings = False
    
    if 'show_memory_browser' not in st.session_state:
        st.session_state.show_memory_browser = False
    
    if 'chat_search' not in st.session_state:
        st.session_state.chat_search = ""


def display_chat_history():
    """
    Display all messages in the chat history with LaTeX and enhanced markdown support.
    Also handles special content like interactive quizzes.
    """
    for message in st.session_state.messages:
        with st.chat_message(message['role']):
            content = message['content']
            
            # Check if this message contains a quiz
            if 'QUIZ_CREATED:' in content:
                # Extract quiz ID and render it
                import re
                quiz_match = re.search(r'QUIZ_CREATED:([^\|]+)\|([^\|]+)\|(.+)', content)
                if quiz_match:
                    quiz_id = quiz_match.group(1)
                    quiz_title = quiz_match.group(2)
                    quiz_info = quiz_match.group(3)
                    
                    # Show a nice message and render the quiz
                    st.success(f"🎓 Quiz ready: **{quiz_title}** ({quiz_info})")
                    render_quiz(quiz_id)
                else:
                    # Fallback to normal rendering
                    render_message_with_latex(content)
            else:
                # Use LaTeX-aware rendering
                render_message_with_latex(content)


@st.dialog("Settings", width="small")
def settings_dialog():
    """Settings dialog - compact."""
    context = get_universal_context()
    identity = context.get('identity', {})
    env = context.get('env', {})
    
    # User + Date inline
    user_name = identity.get('user_name', 'User')
    st.caption(f"👤 {user_name} · {env.get('current_date', '')}")
    
    # AI Toggle
    show_thinking = st.toggle(
        "Show AI reasoning",
        value=st.session_state.get('show_thinking', True),
        key="dialog_thinking"
    )
    st.session_state.show_thinking = show_thinking
    
    st.divider()
    
    # Status - inline compact
    st.caption("**Status**")
    
    # Build status line
    groq_ok = "✅" if GROQ_API_KEY else "❌"
    gmail_ok = "✅" if st.session_state.gmail_authenticated else "⚪"
    
    try:
        collection = get_chroma_collection()
        mem_count = collection.count() if collection else 0
    except:
        mem_count = 0
    
    st.caption(f"Groq {groq_ok} · Gmail {gmail_ok} · Memory: {mem_count}")
    
    # Gmail connect if needed
    if not st.session_state.gmail_authenticated:
        if st.button("Connect Gmail", key="gmail_connect", use_container_width=True):
            try:
                get_gmail_service()
                st.session_state.gmail_authenticated = True
                st.rerun()
            except Exception as e:
                st.error(str(e)[:30])
    else:
        # Gmail is connected - show management buttons
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 Reconnect", key="gmail_reconnect", use_container_width=True):
                result = reconnect_gmail()
                if result is True:
                    st.success("Reconnected!")
                    st.rerun()
                else:
                    st.error(f"Failed: {result}")
        with col2:
            if st.button("❌ Disconnect", key="gmail_disconnect", use_container_width=True):
                if disconnect_gmail():
                    st.success("Disconnected")
                    st.rerun()
        
        # Test connection button
        if st.button("🧪 Test Connection", key="gmail_test", use_container_width=True):
            success, msg = test_gmail_connection()
            if success:
                st.success(msg)
            else:
                st.error(msg)
    
    st.divider()
    
    # Danger zone
    if st.button("🗑️ Clear All Memory", type="secondary", use_container_width=True):
        try:
            import shutil
            global _chroma_client, _chroma_collection
            
            # Reset the global ChromaDB client and collection first
            with _chroma_lock:
                _chroma_client = None
                _chroma_collection = None
            
            # Clear session state
            if 'chroma_collection' in st.session_state:
                del st.session_state.chroma_collection
            
            # Now delete the folder
            if os.path.exists(CHROMA_PATH):
                shutil.rmtree(CHROMA_PATH)
            
            st.success("✅ Memory cleared!")
            st.rerun()
        except Exception as e:
            st.error(f"Error: {e}")


@st.dialog("Memory", width="large")
def memory_dialog():
    """Memory browser - compact."""
    try:
        collection = get_chroma_collection()
        if collection and collection.count() > 0:
            all_docs = collection.get(include=['metadatas', 'documents'])
            
            if all_docs['ids']:
                st.caption(f"{len(all_docs['ids'])} memories stored")
                
                for i, doc_id in enumerate(all_docs['ids'][:10]):
                    meta = all_docs['metadatas'][i] if all_docs['metadatas'] else {}
                    doc_text = all_docs['documents'][i][:150] if all_docs['documents'] else ""
                    subject = meta.get('subject', meta.get('summary', doc_text[:50]))
                    doc_type = meta.get('type', 'doc')
                    
                    with st.expander(f"**{subject[:40]}...** `{doc_type}`" if len(subject) > 40 else f"**{subject}** `{doc_type}`"):
                        st.text(doc_text[:300] + "..." if len(doc_text) > 300 else doc_text)
                
                if len(all_docs['ids']) > 10:
                    st.caption(f"+{len(all_docs['ids']) - 10} more")
        else:
            st.caption("No memories yet")
    except:
        st.caption("Could not load")


@st.dialog("Tasks", width="large")
def tasks_dialog():
    """Task dialog - full width with complete content."""
    task_manager = get_task_manager()
    pending = task_manager.get_pending_tasks()
    done = [t for t in task_manager.get_all_tasks() if t['status'] == 'done']
    
    # Quick add
    new_task = st.text_input("", placeholder="Add new task...", key="dialog_new_task", label_visibility="collapsed")
    if new_task:
        task_manager.add_task(new_task)
        st.rerun()
    
    st.divider()
    
    # Pending tasks - full content
    if pending:
        st.markdown(f"**Pending ({len(pending)})**")
        for task in pending:
            c1, c2 = st.columns([10, 1])
            with c1:
                # Show full task content
                st.markdown(f"• {task['task']}")
            with c2:
                if st.button("✓", key=f"c_{task['id']}", help="Mark done"):
                    task_manager.complete_task_by_id(task['id'])
                    st.rerun()
    else:
        st.caption("No pending tasks")
    
    # Completed - collapsible with full content
    if done:
        st.divider()
        with st.expander(f"Completed ({len(done)})"):
            for t in done[-5:]:
                st.caption(f"- {t['task']}")
            if len(done) > 5:
                st.caption(f"... and {len(done) - 5} more")
            if st.button("Clear completed", key="clear_done", use_container_width=True):
                task_manager.delete_completed()
                st.rerun()


def main():
    """
    Main Streamlit application - Claude-like UI.
    """
    # Page configuration
    st.set_page_config(
        page_title="Emo",
        page_icon="◯",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Apply custom CSS
    st.markdown(get_custom_css(), unsafe_allow_html=True)
    
    # Initialize session state
    initialize_session_state()
    
    # Get context data
    context = get_universal_context()
    identity = context.get('identity', {})
    
    # =========================
    # SIDEBAR - Compact
    # =========================
    with st.sidebar:
        # Header row
        c1, c2 = st.columns([5, 1])
        with c1:
            st.markdown("**Emo**")
        with c2:
            if st.button("＋", key="new_chat_btn", help="New chat"):
                new_id = create_new_session()
                st.session_state.current_session_id = new_id
                st.session_state.messages = []
                set_current_session_id(new_id)
                if 'chat_session' in st.session_state:
                    del st.session_state.chat_session
                st.rerun()
        
        # Search
        chat_search = st.text_input("", placeholder="Search...", key="search_input", label_visibility="collapsed")
        
        # Chats
        st.markdown('<p class="section-label">Chats</p>', unsafe_allow_html=True)
        
        sessions = get_sessions_sorted()
        
        if sessions:
            # Filter by search
            if chat_search:
                sessions = [(sid, sdata) for sid, sdata in sessions 
                           if chat_search.lower() in sdata.get("title", "").lower()]
            
            for session_id, session_data in sessions[:15]:
                title = session_data.get("title", "New Chat")
                is_current = session_id == st.session_state.current_session_id
                
                display_title = title[:35] + "..." if len(title) > 35 else title
                
                col1, col2 = st.columns([6, 1])
                with col1:
                    btn_type = "primary" if is_current else "secondary"
                    if st.button(display_title, key=f"chat_{session_id}", use_container_width=True, type=btn_type):
                        st.session_state.current_session_id = session_id
                        st.session_state.messages = session_data.get("messages", [])
                        set_current_session_id(session_id)
                        if 'chat_session' in st.session_state:
                            del st.session_state.chat_session
                        st.rerun()
                with col2:
                    if st.button("×", key=f"del_{session_id}"):
                        clear_session_short_term_memory(session_id)
                        delete_session(session_id)
                        if session_id == st.session_state.current_session_id:
                            st.session_state.current_session_id = None
                            st.session_state.messages = []
                        st.rerun()
        
        # Bottom section - compact
        st.markdown('<div style="position:fixed;bottom:0;width:inherit;background:#1a1a1a;padding:0.5rem 0;border-top:1px solid #2a2a2a;">', unsafe_allow_html=True)
        
        # Inline action buttons
        task_manager = get_task_manager()
        task_count = len(task_manager.get_pending_tasks())
        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("📋", key="tasks_btn", help=f"Tasks ({task_count})" if task_count else "Tasks"):
                tasks_dialog()
        with c2:
            if st.button("🧠", key="memory_btn", help="Memory"):
                memory_dialog()
        with c3:
            if st.button("⚙️", key="settings_btn", help="Settings"):
                settings_dialog()
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # =========================
    # MAIN CHAT AREA
    # =========================
    
    # Display chat history
    display_chat_history()
    
    # Chat input
    if prompt := st.chat_input("Message Emo..."):
        # Create session if none exists
        if not st.session_state.current_session_id:
            new_id = create_new_session()
            st.session_state.current_session_id = new_id
            set_current_session_id(new_id)
        
        set_current_session_id(st.session_state.current_session_id)
        
        # Add user message
        st.session_state.messages.append({
            'role': 'user',
            'content': prompt
        })
        
        # Generate title from first message
        if len(st.session_state.messages) == 1:
            generate_title(st.session_state.current_session_id, prompt)
        
        # Display user message
        with st.chat_message('user'):
            render_message_with_latex(prompt)
        
        # Generate response
        with st.chat_message('assistant'):
            # Query memory (silent)
            memories = query_memory(prompt)
            memory_context = format_memories_for_context(memories)
            
            # Tool display container
            tool_container = st.container()
            
            # Get response
            show_thinking = st.session_state.get('show_thinking', True)
            result = chat_with_emo(prompt, memory_context, show_thinking, tool_container)
            
            # Display AI thinking in collapsible section (like Gemini/Claude)
            thinking_text = result.get('thinking', '')
            if thinking_text and show_thinking:
                with st.expander("💭 **Thinking**", expanded=False):
                    st.markdown(f"""
<div style="font-size: 0.9rem; color: #8b949e; line-height: 1.6; padding: 0.5rem; background: rgba(139, 148, 158, 0.1); border-radius: 6px; border-left: 3px solid #58a6ff;">
{thinking_text}
</div>
""", unsafe_allow_html=True)
            
            # Display response with LaTeX support
            response_text = result.get('response', '')
            if response_text:
                # Check if this is a quiz response
                if 'QUIZ_CREATED:' in response_text:
                    import re
                    quiz_match = re.search(r'QUIZ_CREATED:([^\|]+)\|([^\|]+)\|(.+)', response_text)
                    if quiz_match:
                        quiz_id = quiz_match.group(1)
                        quiz_title = quiz_match.group(2)
                        quiz_info = quiz_match.group(3)
                        
                        # Show quiz immediately!
                        st.success(f"🎓 Quiz ready: **{quiz_title}** ({quiz_info})")
                        render_quiz(quiz_id)
                    else:
                        render_message_with_latex(response_text)
                else:
                    # Use placeholder for streaming, then render with LaTeX support
                    response_placeholder = st.empty()
                    full_response = ""
                    for word in stream_data(response_text):
                        full_response += word
                        # Live update during streaming
                        response_placeholder.markdown(full_response + "▌")
                    # Final render with proper LaTeX support
                    response_placeholder.empty()
                    render_message_with_latex(full_response)
            response = response_text
        
        # Save to history (without thinking)
        st.session_state.messages.append({
            'role': 'assistant',
            'content': response
        })
        
        if st.session_state.current_session_id:
            save_session(st.session_state.current_session_id, st.session_state.messages)
        
        # Refresh if todos changed
        if st.session_state.get('todo_changed', False):
            st.session_state.todo_changed = False
            st.rerun()


if __name__ == "__main__":
    main()
